"""
ACTA — AI-Powered Research Assistant
============================================
A Streamlit web application for scientific question synthesis
and academic paper framework design.

Dependencies: streamlit, openai, python-docx
"""

import warnings
warnings.filterwarnings("ignore", message=".*numexpr.*")
warnings.filterwarnings("ignore", message=".*bottleneck.*")
import streamlit as st
import streamlit.components.v1 as components
import os
import re
import difflib
import base64
import json
import math
from openai import OpenAI
from docx import Document
from openpyxl import load_workbook
from msoffcrypto.format.ooxml import OOXMLFile
import msoffcrypto
import io
import subprocess

# ==================== Page Configuration ====================
st.set_page_config(
    page_title="ACTA · Research Assistant",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ==================== Custom CSS ====================
st.markdown(
    """
<style>
    /* --- Global: minimal slate-blue palette --- */
    .stApp {
        background-color: #f8fafc;
    }

    .main .block-container {
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
        max-width: 960px !important;
    }

    /* --- Sidebar --- */
    section[data-testid="stSidebar"] {
        background-color: #f1f5f9;
        border-right: 1px solid #e2e8f0;
    }

    section[data-testid="stSidebar"] .block-container {
        padding: 1.5rem 1.2rem;
    }

    /* Sidebar radio buttons */
    div[role="radiogroup"] label {
        padding: 10px 14px !important;
        border-radius: 6px !important;
        margin-bottom: 6px !important;
        font-size: 15px !important;
        font-weight: 500;
        transition: background 0.15s;
    }

    div[role="radiogroup"] label:hover {
        background: #e2e8f0 !important;
    }

    /* Hide radio indicators */
    div[role="radiogroup"] label > div:first-child {
        display: none !important;
    }
    div[role="radiogroup"] label {
        padding-left: 14px !important;
    }

    /* --- Hero-style page titles --- */
    .page-hero {
        font-size: 36px;
        font-weight: 800;
        color: #0f172a;
        letter-spacing: -0.02em;
        margin-bottom: 4px;
        line-height: 1.2;
    }

    .page-tagline {
        font-size: 16px;
        color: #64748b;
        font-weight: 400;
        margin-bottom: 32px;
    }

    /* --- Input fields --- */
    .stTextInput > div > div > input,
    .stTextArea textarea {
        border-radius: 6px !important;
        border: 1px solid #cbd5e1 !important;
        padding: 10px 14px !important;
        font-size: 15px !important;
        background: #ffffff !important;
    }

    .stTextInput > div > div > input:focus,
    .stTextArea textarea:focus {
        border-color: #3b82f6 !important;
        box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.12) !important;
    }

    .stSelectbox > div > div {
        border-radius: 6px !important;
        border: 1px solid #cbd5e1 !important;
    }

    /* --- Buttons --- */
    .stButton > button {
        border-radius: 6px !important;
        font-weight: 600 !important;
        padding: 12px 28px !important;
        border: none !important;
        background: #3b82f6 !important;
        color: #ffffff !important;
        font-size: 15px !important;
        transition: background 0.15s, box-shadow 0.15s !important;
        letter-spacing: 0.01em;
    }

    .stButton > button:hover {
        background: #2563eb !important;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.28) !important;
    }

    .stButton > button:active {
        background: #1d4ed8 !important;
    }

    /* --- Spinner --- */
    .stSpinner > div {
        border-top-color: #3b82f6 !important;
    }

    /* --- Alerts --- */
    div[data-testid="stAlert"] {
        border-radius: 6px;
        padding: 14px 18px;
        font-size: 14px;
    }

    /* --- Dividers --- */
    hr {
        border: none;
        border-top: 1px solid #e2e8f0;
        margin: 32px 0;
    }

    /* --- Question blocks (page 1) --- */
    .question-block {
        background: #f8fafc;
        border-left: 4px solid #3b82f6;
        border-radius: 0 6px 6px 0;
        padding: 18px 22px;
        margin: 16px 0;
    }

    .question-block .q-title {
        font-size: 17px;
        font-weight: 600;
        color: #1e293b;
        line-height: 1.5;
        margin-bottom: 8px;
    }

    .question-block .q-keywords {
        font-size: 13px;
        color: #64748b;
        line-height: 1.6;
    }

    .kw-tag {
        display: inline-block;
        background: #dbeafe;
        color: #1e40af;
        font-size: 12px;
        font-weight: 500;
        border-radius: 4px;
        padding: 2px 8px;
        margin: 2px 4px 2px 0;
    }

    .kw-tag.extended {
        background: #fef3c7;
        color: #92400e;
    }

    /* --- Success banner --- */
    .stSuccess {
        background: #f0fdf4;
        border-left: 4px solid #22c55e;
    }
</style>
""",
    unsafe_allow_html=True,
)

# ==================== Constants ====================
APP_DIR = os.path.dirname(os.path.abspath(__file__))
PROMPT_FILE_PAGE1 = os.path.join(APP_DIR, "prompts_page1.docx")
PROMPT_FILE_PAGE2 = os.path.join(APP_DIR, "prompts_page2.docx")
PROMPT_FILE_PAGE3 = os.path.join(APP_DIR, "prompts_page3.docx")
PROMPT_FILE_PAGE4 = os.path.join(APP_DIR, "prompts_page4.docx")
PROMPT_FILE_PAGE5 = os.path.join(APP_DIR, "prompts_page5.docx")
PET_DIR = os.path.join(APP_DIR, "desktop_pet")
PET_SPRITESHEET = os.path.join(PET_DIR, "spritesheet.webp")
PET_VALIDATION = os.path.join(PET_DIR, "validation.json")
PET_LAYOUT_TXT = os.path.join(PET_DIR, "web_pet_layout.txt")
PET_LAYOUT_XLSX = os.path.join(PET_DIR, "web_pet_layout.xlsx")
ACTA_URL = "https://actaresearch.streamlit.app/"


WEB_PET_LAYOUT_DEFAULTS = {
    "frame_width": 440,
    "frame_height": 450,
    "frame_right": 24,
    "frame_bottom": 24,
    "pet_left": 176,
    "pet_top": 280,
    "pet_scale": 0.8,
    "pet_base_width": 160,
    "pet_base_height": 176,
    "speech_right": 28,
    "speech_top": 112,
    "speech_min_width": 190,
    "speech_max_width": 320,
    "qa_left": 16,
    "qa_top": 72,
    "qa_initial_width": 230,
    "qa_final_width": 230,
    "qa_width": 230,
    "qa_log_height": 128,
    "qa_padding": 10,
    "qa_font_size": 12,
    "qa_button_width": 34,
    "qa_gap": 6,
    "quip_interval_ms": 20000,
    "quip_request_timeout_ms": 8000,
    "quip_fallback_enabled": 1,
}


def _coerce_layout_value(value):
    """Convert spreadsheet/text layout values to numbers when possible."""
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return int(value) if float(value).is_integer() else float(value)
    text = str(value).strip()
    if not text:
        return None
    if text.lower().endswith("px"):
        text = text[:-2].strip()
    try:
        number = float(text)
        return int(number) if number.is_integer() else number
    except ValueError:
        return text


def load_web_pet_layout():
    """Load adjustable web pet layout values from xlsx or text config."""
    layout = dict(WEB_PET_LAYOUT_DEFAULTS)

    def apply_value(key, value):
        key = str(key or "").strip()
        if key not in layout:
            return
        parsed = _coerce_layout_value(value)
        if parsed is not None:
            layout[key] = parsed

    try:
        if os.path.exists(PET_LAYOUT_XLSX):
            workbook = load_workbook(PET_LAYOUT_XLSX, data_only=True)
            sheet = workbook.active
            for row in sheet.iter_rows(min_row=1, values_only=True):
                if not row or len(row) < 2:
                    continue
                apply_value(row[0], row[1])
            return layout
    except Exception:
        pass

    try:
        if os.path.exists(PET_LAYOUT_TXT):
            with open(PET_LAYOUT_TXT, "r", encoding="utf-8") as layout_file:
                for raw_line in layout_file:
                    line = raw_line.strip()
                    if not line or line.startswith("#"):
                        continue
                    if "=" in line:
                        key, value = line.split("=", 1)
                    elif "," in line:
                        key, value = line.split(",", 1)
                    else:
                        continue
                    apply_value(key, value)
    except Exception:
        pass

    return layout


def render_web_pet():
    """Render the desktop pet as a floating in-browser companion."""
    if st.session_state.get("web_pet_disabled"):
        return
    if not os.path.exists(PET_SPRITESHEET) or not os.path.exists(PET_VALIDATION):
        return

    try:
        with open(PET_SPRITESHEET, "rb") as sprite_file:
            sprite_data = base64.b64encode(sprite_file.read()).decode("ascii")
        with open(PET_VALIDATION, "r", encoding="utf-8") as validation_file:
            metadata = json.load(validation_file)
    except Exception:
        return

    pet_config = {
        "sprite": "data:image/webp;base64," + sprite_data,
        "metadata": metadata,
        "layout": load_web_pet_layout(),
        "actaUrl": ACTA_URL,
        "apiKey": st.session_state.get("api_key", "").strip(),
        "baseUrl": st.session_state.get("base_url", "https://api.siliconflow.cn/v1").strip(),
        "model": st.session_state.get("model", "deepseek-ai/DeepSeek-V4-Pro").strip(),
        "apiEnabled": bool(
            st.session_state.get("api_key", "").strip()
            or st.session_state.get("user_key_valid")
        ),
    }

    html = """
<script>
(function () {
  const config = __PET_CONFIG__;
  const layout = Object.assign({
    frame_width: 440,
    frame_height: 450,
    frame_right: 24,
    frame_bottom: 24,
    pet_left: 176,
    pet_top: 280,
    pet_scale: 0.8,
    pet_base_width: 160,
    pet_base_height: 176,
    speech_right: 28,
    speech_top: 112,
    speech_min_width: 190,
    speech_max_width: 320,
    qa_left: 16,
    qa_top: 72,
    qa_initial_width: 230,
    qa_final_width: 230,
    qa_width: 230,
    qa_log_height: 128,
    qa_padding: 10,
    qa_font_size: 12,
    qa_button_width: 34,
    qa_gap: 6,
    quip_interval_ms: 20000,
    quip_request_timeout_ms: 8000,
    quip_fallback_enabled: 1,
  }, config.layout || {});
  if (config.layout && config.layout.qa_width != null) {
    if (config.layout.qa_initial_width == null) layout.qa_initial_width = config.layout.qa_width;
    if (config.layout.qa_final_width == null) layout.qa_final_width = config.layout.qa_width;
  }
  function layoutNumber(key, fallback) {
    const value = Number(layout[key]);
    return Number.isFinite(value) ? value : fallback;
  }
  function layoutPx(key, fallback) {
    return `${layoutNumber(key, fallback)}px`;
  }
  const doc = document;
  const frame = window.frameElement;
  if (frame) {
    frame.style.position = "fixed";
    frame.style.right = layoutPx("frame_right", 24);
    frame.style.bottom = layoutPx("frame_bottom", 24);
    frame.style.width = layoutPx("frame_width", 440);
    frame.style.height = layoutPx("frame_height", 450);
    frame.style.border = "0";
    frame.style.zIndex = "2147483000";
    frame.style.background = "transparent";
  }
  doc.documentElement.style.background = "transparent";
  if (doc.body) {
    doc.body.style.margin = "0";
    doc.body.style.overflow = "hidden";
    doc.body.style.background = "transparent";
  }

  const oldRoot = doc.getElementById("acta-web-pet-root");
  if (oldRoot) oldRoot.remove();
  const oldStyle = doc.getElementById("acta-web-pet-style");
  if (oldStyle) oldStyle.remove();

  const style = doc.createElement("style");
  style.id = "acta-web-pet-style";
  style.textContent = `
    #acta-web-pet-root {
      position: absolute;
      left: ${layoutPx("pet_left", 176)};
      top: ${layoutPx("pet_top", 280)};
      z-index: 2147483000;
      width: ${layoutPx("pet_base_width", 160)};
      height: ${layoutPx("pet_base_height", 176)};
      user-select: none;
      touch-action: none;
      pointer-events: auto;
    }
    #acta-web-pet-sprite {
      position: relative;
      width: 100%;
      height: 100%;
      background-repeat: no-repeat;
      cursor: grab;
      filter: drop-shadow(0 10px 18px rgba(15, 23, 42, 0.18));
    }
    #acta-web-pet-effect {
      position: absolute;
      inset: 0;
      pointer-events: none;
      display: none;
    }
    #acta-web-pet-root.acta-state-running #acta-web-pet-sprite {
      animation: acta-paper-bob 520ms steps(2, end) infinite;
    }
    #acta-web-pet-root.acta-state-running #acta-web-pet-effect {
      display: block;
    }
    #acta-web-pet-root.acta-state-running #acta-web-pet-effect::before {
      content: "...";
      position: absolute;
      left: 80px;
      top: 98px;
      width: 52px;
      height: 22px;
      border-radius: 999px;
      color: #2563eb;
      font-size: 22px;
      font-weight: 900;
      letter-spacing: 2px;
      line-height: 1;
      text-shadow: 0 0 7px rgba(96, 165, 250, 0.9), 0 1px 0 rgba(255,255,255,0.9);
      animation: acta-typing-dots 620ms steps(3, end) infinite;
    }
    #acta-web-pet-root.acta-state-running #acta-web-pet-effect::after {
      content: "";
      position: absolute;
      left: 58px;
      top: 86px;
      width: 68px;
      height: 50px;
      border-radius: 8px;
      background: radial-gradient(circle at 50% 35%, rgba(147, 197, 253, 0.36), transparent 62%);
      animation: acta-screen-pulse 900ms ease-in-out infinite;
    }
    #acta-web-pet-root.acta-state-review #acta-web-pet-sprite {
      animation: acta-lab-focus 780ms ease-in-out infinite;
    }
    #acta-web-pet-root.acta-state-review #acta-web-pet-effect {
      display: block;
    }
    #acta-web-pet-root.acta-state-review #acta-web-pet-effect::before {
      content: "o  o  o";
      position: absolute;
      left: 34px;
      top: 88px;
      width: 62px;
      height: 58px;
      color: #0891b2;
      font-size: 22px;
      font-weight: 900;
      line-height: 0.9;
      white-space: pre;
      text-shadow: 0 0 8px rgba(45, 212, 191, 0.9), 0 1px 0 rgba(255,255,255,0.9);
      animation: acta-lab-bubbles 980ms linear infinite;
    }
    #acta-web-pet-root.acta-state-review #acta-web-pet-effect::after {
      content: "";
      position: absolute;
      left: 42px;
      top: 132px;
      width: 38px;
      height: 12px;
      border-radius: 50%;
      background: rgba(45, 212, 191, 0.42);
      box-shadow: 0 0 10px rgba(45, 212, 191, 0.75);
      animation: acta-liquid-pulse 680ms ease-in-out infinite;
    }
    @keyframes acta-paper-bob {
      0%, 100% { transform: translateY(0); }
      50% { transform: translateY(-2px); }
    }
    @keyframes acta-typing-dots {
      0% { clip-path: inset(0 36px 0 0); opacity: 0.50; transform: translateY(2px); }
      33% { clip-path: inset(0 18px 0 0); opacity: 0.78; transform: translateY(0); }
      66%, 100% { clip-path: inset(0 0 0 0); opacity: 1; }
    }
    @keyframes acta-screen-pulse {
      0%, 100% { opacity: 0.22; transform: scale(0.96); }
      50% { opacity: 0.72; transform: scale(1.06); }
    }
    @keyframes acta-lab-focus {
      0%, 100% { transform: rotate(0deg) translateY(0); }
      50% { transform: rotate(-1.5deg) translateY(-1px); }
    }
    @keyframes acta-lab-bubbles {
      0% { opacity: 0.25; transform: translateY(8px); }
      45% { opacity: 1; }
      100% { opacity: 0; transform: translateY(-12px); }
    }
    @keyframes acta-liquid-pulse {
      0%, 100% { opacity: 0.35; transform: scaleX(0.85); }
      50% { opacity: 0.95; transform: scaleX(1.18); }
    }
    #acta-web-pet-root.acta-state-running-left #acta-web-pet-sprite,
    #acta-web-pet-root.acta-state-running-right #acta-web-pet-sprite {
      animation: acta-drag-run 320ms steps(2, end) infinite;
    }
    @keyframes acta-drag-run {
      0%, 100% { transform: translateY(0) rotate(0deg); }
      50% { transform: translateY(-3px) rotate(-1deg); }
    }
    #acta-web-pet-root.dragging #acta-web-pet-sprite {
      cursor: grabbing;
    }
    #acta-web-pet-countdown {
      position: absolute;
      left: 50%;
      top: -34px;
      transform: translateX(-50%);
      min-width: 96px;
      padding: 5px 10px;
      border: 1px solid rgba(59, 130, 246, 0.25);
      border-radius: 999px;
      background: rgba(255, 255, 255, 0.96);
      box-shadow: 0 8px 18px rgba(15, 23, 42, 0.14);
      color: #0f172a;
      display: none;
      font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      font-size: 12px;
      font-weight: 700;
      line-height: 1;
      text-align: center;
      white-space: nowrap;
    }
    #acta-web-pet-countdown.show {
      display: block;
    }
    #acta-web-pet-speech {
      position: fixed;
      right: ${layoutPx("speech_right", 28)};
      top: ${layoutPx("speech_top", 112)};
      transform: none;
      width: max-content;
      min-width: ${layoutPx("speech_min_width", 190)};
      max-width: ${layoutPx("speech_max_width", 320)};
      padding: 11px 14px;
      border: 1px solid rgba(14, 165, 233, 0.22);
      border-radius: 12px;
      background: rgba(255, 255, 255, 0.97);
      box-shadow: 0 10px 22px rgba(15, 23, 42, 0.14);
      color: #0f172a;
      display: none;
      font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      font-size: 14px;
      font-weight: 650;
      line-height: 1.42;
      text-align: left;
      pointer-events: none;
      z-index: 3;
    }
    #acta-web-pet-speech::after {
      content: "";
      position: absolute;
      right: 28px;
      bottom: -6px;
      width: 10px;
      height: 10px;
      background: rgba(255, 255, 255, 0.97);
      border-right: 1px solid rgba(14, 165, 233, 0.18);
      border-bottom: 1px solid rgba(14, 165, 233, 0.18);
      transform: rotate(45deg);
    }
    #acta-web-pet-speech.show {
      display: block;
      animation: acta-speech-pop 180ms ease-out;
    }
    @keyframes acta-speech-pop {
      from { opacity: 0; transform: translateY(5px) scale(0.96); }
      to { opacity: 1; transform: translateY(0) scale(1); }
    }
    #acta-web-pet-qa {
      position: fixed;
      left: ${layoutPx("qa_left", 16)};
      top: ${layoutPx("qa_top", 72)};
      transform: none;
      width: ${layoutPx("qa_initial_width", 230)};
      min-width: 0;
      max-width: ${layoutPx("qa_initial_width", 230)};
      box-sizing: border-box;
      overflow: hidden;
      padding: ${layoutPx("qa_padding", 10)};
      border: 1px solid rgba(37, 99, 235, 0.65);
      border-radius: 8px;
      background: #eff6ff;
      box-shadow: none;
      color: #0f172a;
      display: none;
      font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      z-index: 4;
    }
    #acta-web-pet-qa.open {
      display: grid;
      gap: 8px;
      width: ${layoutPx("qa_final_width", 230)};
      max-width: ${layoutPx("qa_final_width", 230)};
    }
    #acta-web-pet-qa-log {
      max-height: ${layoutPx("qa_log_height", 128)};
      overflow: auto;
      padding: 8px;
      border-radius: 6px;
      background: #f8fafc;
      font-size: ${layoutPx("qa_font_size", 12)};
      line-height: 1.45;
      overflow-wrap: anywhere;
      word-break: break-word;
    }
    #acta-web-pet-qa-log div {
      margin-bottom: 6px;
    }
    #acta-web-pet-qa-log b {
      color: #1d4ed8;
    }
    #acta-web-pet-qa-row {
      display: grid;
      grid-template-columns: minmax(0, 1fr) ${layoutPx("qa_button_width", 34)} ${layoutPx("qa_button_width", 34)};
      gap: ${layoutPx("qa_gap", 6)};
      min-width: 0;
    }
    #acta-web-pet-qa input {
      min-width: 0;
      border: 1px solid #cbd5e1;
      border-radius: 6px;
      padding: 7px 8px;
      font-size: ${layoutPx("qa_font_size", 12)};
      outline: none;
    }
    #acta-web-pet-qa input:focus {
      border-color: #3b82f6;
      box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.12);
    }
    #acta-web-pet-qa button {
      border: 0;
      border-radius: 6px;
      background: #dbeafe;
      color: #1e3a8a;
      font-size: 13px;
      font-weight: 800;
      cursor: pointer;
    }
    #acta-web-pet-qa button:hover {
      background: #bfdbfe;
    }
    #acta-web-pet-menu {
      position: absolute;
      min-width: 174px;
      padding: 8px;
      border: 1px solid rgba(148, 163, 184, 0.45);
      border-radius: 8px;
      background: rgba(255, 255, 255, 0.98);
      box-shadow: 0 16px 35px rgba(15, 23, 42, 0.20);
      display: none;
      font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    }
    #acta-web-pet-menu.open {
      display: grid;
      gap: 5px;
    }
    #acta-web-pet-menu button {
      width: 100%;
      min-height: 30px;
      border: 0;
      border-radius: 6px;
      background: #f1f5f9;
      color: #0f172a;
      font-size: 13px;
      text-align: left;
      padding: 6px 9px;
      cursor: pointer;
    }
    #acta-web-pet-menu button:hover {
      background: #dbeafe;
    }
    #acta-web-pet-menu [hidden] {
      display: none !important;
    }
    #acta-web-pet-menu .row {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 5px;
    }
    #acta-web-pet-menu .single {
      grid-template-columns: 1fr;
    }
    #acta-web-pet-menu .divider {
      height: 1px;
      background: #e2e8f0;
      margin: 3px 0;
    }
    #acta-web-pet-menu .menu-title {
      color: #64748b;
      font-size: 12px;
      font-weight: 700;
      padding: 2px 2px 4px;
    }
    @media (max-width: 640px) {
      #acta-web-pet-root {
        left: ${layoutPx("pet_left", 176)};
        top: ${layoutPx("pet_top", 280)};
      }
      #acta-web-pet-menu {
        right: 0;
        min-width: 160px;
      }
    }
  `;
  (doc.head || doc.documentElement).appendChild(style);

  const root = doc.createElement("div");
  root.id = "acta-web-pet-root";
  root.innerHTML = `
    <div id="acta-web-pet-speech" aria-live="polite"></div>
    <div id="acta-web-pet-qa" aria-live="polite">
      <div id="acta-web-pet-qa-log">
        <div><b>ACTA</b>: &#25105;&#20250;&#25351;&#24341;&#20320;&#21435;&#23545;&#24212;&#30340;&#32593;&#31449;&#21151;&#33021;&#65292;&#19981;&#30452;&#25509;&#26367;&#20195;&#20320;&#23436;&#25104;&#30740;&#31350;&#20219;&#21153;&#12290;</div>
      </div>
      <div id="acta-web-pet-qa-row">
        <input id="acta-web-pet-qa-input" type="text" maxlength="180" placeholder="&#31616;&#30701;&#38382;&#19968;&#21477;..." />
        <button id="acta-web-pet-qa-send" title="Send">&#8594;</button>
        <button id="acta-web-pet-qa-close" title="Close">&#215;</button>
      </div>
    </div>
    <div id="acta-web-pet-countdown" aria-live="polite"></div>
    <div id="acta-web-pet-sprite" title="ACTA Labmate"></div>
    <div id="acta-web-pet-effect"></div>
    <div id="acta-web-pet-menu" aria-label="ACTA Labmate menu">
      <div data-menu-panel="main">
        <div class="row">
          <button data-action="idle">&#24453;&#26426;</button>
          <button data-action="think">&#24605;&#32771;</button>
        </div>
        <div class="divider"></div>
        <div class="row single">
          <button data-timed-action="rest">&#20241;&#24687; &#8250;</button>
          <button data-timed-action="paper">&#35770;&#25991; &#8250;</button>
          <button data-timed-action="experiment">&#23454;&#39564; &#8250;</button>
        </div>
        <div class="divider"></div>
        <div class="row">
          <button id="acta-web-pet-talk-toggle" data-talk-toggle="true">&#39578;&#35805;&#27169;&#24335;</button>
          <button data-qa-mode="true">QA &#27169;&#24335;</button>
        </div>
        <div class="divider"></div>
        <div class="row single">
          <button data-menu-target="size">&#22823;&#23567; &#8250;</button>
        </div>
        <div class="divider"></div>
        <button data-close="true">&#36864;&#20986;</button>
      </div>
      <div data-menu-panel="duration" hidden>
        <div class="menu-title" data-duration-title></div>
        <div class="row single">
          <button data-duration="300">5 &#20998;&#38047;</button>
          <button data-duration="1800">30 &#20998;&#38047;</button>
          <button data-duration="3600">1 &#23567;&#26102;</button>
          <button data-duration="7200">2 &#23567;&#26102;</button>
        </div>
        <div class="divider"></div>
        <button data-menu-target="main">&#8249; &#36820;&#22238;</button>
      </div>
      <div data-menu-panel="size" hidden>
        <div class="menu-title">&#22823;&#23567;</div>
        <div class="row">
          <button data-size="0.6">60%</button>
          <button data-size="0.8">80%</button>
        </div>
        <div class="row">
          <button data-size="1">100%</button>
          <button data-size="1.2">120%</button>
        </div>
        <div class="divider"></div>
        <button data-menu-target="main">&#8249; &#36820;&#22238;</button>
      </div>
    </div>
  `;
  (doc.body || doc.documentElement).appendChild(root);

  const sprite = root.querySelector("#acta-web-pet-sprite");
  const speech = root.querySelector("#acta-web-pet-speech");
  const qa = root.querySelector("#acta-web-pet-qa");
  const qaLog = root.querySelector("#acta-web-pet-qa-log");
  const qaInput = root.querySelector("#acta-web-pet-qa-input");
  const qaSend = root.querySelector("#acta-web-pet-qa-send");
  const qaClose = root.querySelector("#acta-web-pet-qa-close");
  const countdown = root.querySelector("#acta-web-pet-countdown");
  const menu = root.querySelector("#acta-web-pet-menu");
  const talkToggleButton = root.querySelector("#acta-web-pet-talk-toggle");
  const panels = {
    main: root.querySelector('[data-menu-panel="main"]'),
    duration: root.querySelector('[data-menu-panel="duration"]'),
    size: root.querySelector('[data-menu-panel="size"]'),
  };
  const durationTitle = root.querySelector("[data-duration-title]");
  const meta = config.metadata;
  const cellWidth = meta.width / meta.columns;
  const cellHeight = meta.height / meta.rows;
  const groups = {};

  for (const cell of meta.cells || []) {
    if (!cell.used) continue;
    if (!groups[cell.state]) groups[cell.state] = [];
    groups[cell.state].push({ row: cell.row, column: cell.column });
  }
  for (const state of Object.keys(groups)) {
    groups[state].sort((a, b) => (a.row - b.row) || (a.column - b.column));
  }

  const actionStates = {
    idle: "idle",
    rest: "waiting",
    paper: "running",
    experiment: "review",
    think: "thinking",
    wave: "waving",
    jump: "jumping",
  };
  const timedActions = new Set(["rest", "paper", "experiment"]);
  const actionQuips = {
    idle: [
      "先让变量冷静一下。",
      "今日份科研玄学加载中。",
      "空白不是偷懒，是在预处理灵感。",
    ],
    thinking: [
      "这个假设有点东西。",
      "脑内 peer review 进行中。",
      "先别急，我在和 p 值谈判。",
    ],
    waiting: [
      "咖啡是液态置信区间。",
      "休息一下，方法学会自己长出来。",
      "把大脑送进低功耗模式。",
    ],
    running: [
      "键盘正在产出高影响因子噪声。",
      "摘要先跑，结果随后就到。",
      "正在把灵感编译成论文。",
    ],
    review: [
      "试管里冒的是显著性。",
      "实验组正在努力不离谱。",
      "样品说它也想上一区。",
    ],
    waving: [
      "导师路过，保持微笑。",
      "向审稿人释放友好信号。",
    ],
    jumping: [
      "灵感出现，先跳为敬。",
      "这个 idea 值得开新坑。",
    ],
    "running-left": [
      "样本向左漂移，正在追踪。",
      "别让灵感逃出视窗。",
    ],
    "running-right": [
      "向右迁移，保持梯度稳定。",
      "追上鼠标，追上 deadline。",
    ],
  };
  const actionLabels = {
    idle: "待机",
    thinking: "思考",
    waiting: "休息",
    running: "写论文",
    review: "做实验",
    waving: "刚进入网站挥手",
    jumping: "点击跳跃",
    "running-left": "向左拖动",
    "running-right": "向右拖动",
  };
  let state = "";
  let returnState = "idle";
  let oneShotState = null;
  let frameIndex = 0;
  let scale = layoutNumber("pet_scale", Number(localStorage.getItem("actaPetScale") || "0.8"));
  let timer = null;
  let countdownTimer = null;
  let countdownEndAt = 0;
  let clickTimer = null;
  let selectedTimedAction = null;
  let speechTimer = null;
  let talkTimer = null;
  let talkMode = localStorage.getItem("actaPetTalkMode") === "true";
  let qaMode = false;
  let llmBusy = false;
  let lastSpeechAt = 0;
  let pointerDown = false;
  let dragging = false;
  let pressButton = 0;
  let offsetX = 0;
  let offsetY = 0;
  let framePointerOffsetX = 0;
  let framePointerOffsetY = 0;
  let pressX = 0;
  let pressY = 0;
  let lastPointerX = 0;
  let lastScreenX = 0;
  let lastScreenY = 0;
  let dragState = "running-right";

  function clamp(value, min, max) {
    return Math.min(Math.max(value, min), max);
  }

  function hostViewport() {
    try {
      return {
        width: window.parent.innerWidth || window.innerWidth,
        height: window.parent.innerHeight || window.innerHeight,
      };
    } catch (error) {
      return { width: window.innerWidth, height: window.innerHeight };
    }
  }

  function parentInputValue(selector) {
    try {
      const parentDoc = window.parent.document;
      const input = parentDoc.querySelector(selector);
      return input ? String(input.value || "").trim() : "";
    } catch (error) {
      return "";
    }
  }

  function llmApiKey() {
    return String(config.apiKey || "").trim() || parentInputValue('input[placeholder="sk-..."]');
  }

  function llmBaseUrl() {
    return String(config.baseUrl || "").trim()
      || parentInputValue('input[placeholder="https://api.siliconflow.cn/v1"]')
      || "https://api.siliconflow.cn/v1";
  }

  function llmModel() {
    return String(config.model || "").trim()
      || parentInputValue('input[placeholder="deepseek-ai/DeepSeek-V4-Pro"]')
      || "deepseek-ai/DeepSeek-V4-Pro";
  }

  function apiAvailable() {
    return llmApiKey().length > 0;
  }

  function chatEndpoint() {
    return `${llmBaseUrl().replace(/\/+$/, "")}/chat/completions`;
  }

  async function requestLlm(messages, maxTokens = 80) {
    const apiKey = llmApiKey();
    if (!apiKey) {
      throw new Error("missing api key");
    }
    const controller = new AbortController();
    const timeout = window.setTimeout(
      () => controller.abort(),
      Math.max(1000, layoutNumber("quip_request_timeout_ms", 8000))
    );
    const response = await fetch(chatEndpoint(), {
      method: "POST",
      signal: controller.signal,
      headers: {
        "Content-Type": "application/json",
        "Authorization": `Bearer ${apiKey}`,
      },
      body: JSON.stringify({
        model: llmModel(),
        messages,
        temperature: 0.85,
        max_tokens: maxTokens,
      }),
    }).finally(() => window.clearTimeout(timeout));
    if (!response.ok) {
      throw new Error(`llm http ${response.status}`);
    }
    const data = await response.json();
    return String(data?.choices?.[0]?.message?.content || "").trim();
  }

  function applyScale(value) {
    scale = Number(value) || 0.8;
    localStorage.setItem("actaPetScale", String(scale));
    root.style.width = `${Math.round(cellWidth * scale)}px`;
    root.style.height = `${Math.round(cellHeight * scale)}px`;
    sprite.style.backgroundImage = `url("${config.sprite}")`;
    sprite.style.backgroundSize = `${Math.round(meta.width * scale)}px ${Math.round(meta.height * scale)}px`;
    const rect = root.getBoundingClientRect();
    root.style.left = `${clamp(rect.left, 8, window.innerWidth - root.offsetWidth - 8)}px`;
    root.style.top = `${clamp(rect.top, 8, window.innerHeight - root.offsetHeight - 8)}px`;
  }

  function setState(nextState) {
    const resolvedState = groups[nextState] ? nextState : "idle";
    if (state === resolvedState) {
      return;
    }
    state = resolvedState;
    frameIndex = 0;
    for (const className of Array.from(root.classList)) {
      if (className.startsWith("acta-state-")) {
        root.classList.remove(className);
      }
    }
    root.classList.add(`acta-state-${state}`);
  }

  function hideSpeech() {
    if (speechTimer) {
      window.clearTimeout(speechTimer);
      speechTimer = null;
    }
    speech.classList.remove("show");
  }

  function updateTalkToggleButton() {
    if (!talkToggleButton) return;
    talkToggleButton.textContent = talkMode ? "关闭骚话" : "骚话模式";
    talkToggleButton.title = talkMode ? "关闭 LLM 自动科研骚话" : "开启 LLM 自动科研骚话";
  }

  function showSpeechText(text, duration = 10000) {
    if (!speech || !text) return;
    speech.textContent = String(text).replace(/["“”]/g, "").slice(0, 80);
    speech.classList.add("show");
    if (speechTimer) {
      window.clearTimeout(speechTimer);
    }
    speechTimer = window.setTimeout(() => {
      speech.classList.remove("show");
    }, duration);
  }

  function fallbackQuip(actionState) {
    const choices = {
      idle: ["变量先冷静一下。", "今天也在加载科研玄学。", "空白不是偷懒，是预处理灵感。"],
      thinking: ["这个假设有点东西。", "脑内 peer review 中。", "先别急，我在和 p 值谈判。"],
      waiting: ["休息也是方法学的一部分。", "咖啡是液态置信区间。", "把大脑送进低功耗模式。"],
      running: ["正在把灵感翻译成论文。", "摘要先跑，结果随后就到。", "键盘正在产出影响因子噪声。"],
      review: ["试管里冒的是显著性。", "样品说它也想上一作。", "实验组正在努力不离谱。"],
      "running-left": ["别让灵感逃出窗口。", "样本向左漂移，正在追踪。"],
      "running-right": ["向右迁移，梯度稳定。", "追上鼠标，也追 deadline。"],
    };
    const list = choices[actionState] || choices.idle;
    return list[Math.floor(Math.random() * list.length)];
  }

  async function showActionSpeech(force = false) {
    if (!talkMode || qaMode || !speech || llmBusy) {
      return;
    }
    if (!apiAvailable()) {
      if (force) showSpeechText("请先在左侧 API Configuration 填写 API Key。", 10000);
      return;
    }
    const now = Date.now();
    const quipInterval = Math.max(1000, layoutNumber("quip_interval_ms", 20000));
    if (!force && now - lastSpeechAt < Math.max(0, quipInterval - 500)) {
      return;
    }
    llmBusy = true;
    lastSpeechAt = now;
    try {
      const label = actionLabels[state] || "科研陪伴";
      const text = await requestLlm([
        {
          role: "system",
          content: "你是网页桌宠 ACTA Labmate。根据当前动作生成一句中文科研吐槽或鼓励，幽默、短、自然，不超过22个汉字，不要解释，不要引号。",
        },
        {
          role: "user",
          content: `当前动作：${label}。请生成一句适合飘在角色上方的科研骚话。`,
        },
      ], 48);
      if (!talkMode || qaMode || !text) return;
      speech.textContent = text.replace(/["“”]/g, "").slice(0, 60);
      speech.classList.add("show");
      if (speechTimer) {
        window.clearTimeout(speechTimer);
      }
      speechTimer = window.setTimeout(() => {
        speech.classList.remove("show");
      }, 10000);
    } catch (error) {
      console.warn("ACTA pet speech request failed", error);
      if (layoutNumber("quip_fallback_enabled", 1)) {
        showSpeechText(fallbackQuip(state), 10000);
      }
    } finally {
      llmBusy = false;
    }
  }

  function refreshTalkTimer(runNow = false) {
    if (talkTimer) {
      window.clearInterval(talkTimer);
      talkTimer = null;
    }
    if (!talkMode) {
      hideSpeech();
      return;
    }
    if (!apiAvailable()) {
      if (runNow) showSpeechText("请先在左侧 API Configuration 填写 API Key。", 10000);
      return;
    }
    if (runNow) showActionSpeech(true);
    talkTimer = window.setInterval(
      () => showActionSpeech(false),
      Math.max(1000, layoutNumber("quip_interval_ms", 20000))
    );
  }

  function appendQaLine(name, text) {
    const line = doc.createElement("div");
    const who = doc.createElement("b");
    who.textContent = name;
    line.appendChild(who);
    line.appendChild(doc.createTextNode(`: ${text}`));
    qaLog.appendChild(line);
    qaLog.scrollTop = qaLog.scrollHeight;
  }

  function openQa() {
    qaMode = true;
    hideSpeech();
    qa.classList.add("open");
    if (!apiAvailable()) {
      appendQaLine("ACTA", "我可以先告诉你功能在哪里；要让我用 LLM 回答，请在左侧 API Configuration 填写 API Key。");
    }
    qaInput.focus();
  }

  function closeQa() {
    qaMode = false;
    qa.classList.remove("open");
    refreshTalkTimer(false);
  }

  async function sendQa() {
    const question = qaInput.value.trim();
    if (!question || llmBusy) return;
    qaInput.value = "";
    appendQaLine("你", question);
    appendQaLine("ACTA", "思考中...");
    const pending = qaLog.lastElementChild;
    if (!apiAvailable()) {
      pending.innerHTML = "";
      const who = doc.createElement("b");
      who.textContent = "ACTA";
      pending.appendChild(who);
      pending.appendChild(doc.createTextNode(": 请先到左侧 API Configuration 填写 API Key。功能位置：期刊推荐在 Journal Recommender，科研问题在 Scientific Question Synthesis，框架在 Framework Architect，引言在 Introduction Generator。"));
      qaLog.scrollTop = qaLog.scrollHeight;
      return;
    }
    llmBusy = true;
    try {
      const answer = await requestLlm([
        {
          role: "system",
          content: "你是 ACTA Research Assistant 网站导航桌宠。不要直接完成用户的科研任务，优先判断用户想找什么功能，并指引他到网站位置：Journal Recommender 用于期刊推荐，Scientific Question Synthesis 用于生成科研问题，Framework Architect 用于论文框架，Introduction Generator 用于引言生成，API Configuration 在左侧边栏用于填写 API Key/Base URL/Model，About ACTA 查看说明。中文简短回答，不超过90字。",
        },
        { role: "user", content: question },
      ], 140);
      pending.innerHTML = "";
      const who = doc.createElement("b");
      who.textContent = "ACTA";
      pending.appendChild(who);
      pending.appendChild(doc.createTextNode(`: ${answer || "我刚刚没组织好语言，请再问一次。"}`));
    } catch (error) {
      pending.innerHTML = "";
      const who = doc.createElement("b");
      who.textContent = "ACTA";
      pending.appendChild(who);
      pending.appendChild(doc.createTextNode(": API 请求失败，请检查 Key、Base URL、模型或浏览器跨域限制。"));
      console.warn("ACTA pet QA request failed", error);
    } finally {
      llmBusy = false;
      qaLog.scrollTop = qaLog.scrollHeight;
    }
  }

  function formatRemaining(totalSeconds) {
    const seconds = Math.max(0, Math.ceil(totalSeconds));
    const hours = Math.floor(seconds / 3600);
    const minutes = Math.floor((seconds % 3600) / 60);
    const restSeconds = seconds % 60;
    if (hours > 0) {
      return `${hours}:${String(minutes).padStart(2, "0")}:${String(restSeconds).padStart(2, "0")}`;
    }
    return `${minutes}:${String(restSeconds).padStart(2, "0")}`;
  }

  function updateCountdown() {
    if (!countdownEndAt) {
      countdown.textContent = "";
      countdown.classList.remove("show");
      return;
    }
    const remaining = Math.max(0, (countdownEndAt - Date.now()) / 1000);
    countdown.textContent = formatRemaining(remaining);
    countdown.classList.add("show");
    if (remaining <= 0) {
      finishTimedAction();
    }
  }

  function clearTimer() {
    if (timer) window.clearTimeout(timer);
    timer = null;
    if (countdownTimer) window.clearInterval(countdownTimer);
    countdownTimer = null;
    countdownEndAt = 0;
    updateCountdown();
    oneShotState = null;
  }

  function finishTimedAction() {
    if (timer) window.clearTimeout(timer);
    if (countdownTimer) window.clearInterval(countdownTimer);
    timer = null;
    countdownTimer = null;
    countdownEndAt = 0;
    updateCountdown();
    oneShotState = null;
    returnState = "idle";
    if (!dragging) setState("idle");
  }

  function setFreeAction(action) {
    clearTimer();
    const nextState = actionStates[action] || "idle";
    returnState = nextState;
    setState(nextState);
  }

  function setTimedAction(action, seconds) {
    clearTimer();
    const nextState = actionStates[action] || "idle";
    returnState = nextState;
    setState(nextState);
    countdownEndAt = Date.now() + seconds * 1000;
    updateCountdown();
    countdownTimer = window.setInterval(updateCountdown, 1000);
    timer = window.setTimeout(finishTimedAction, seconds * 1000);
  }

  function playOnce(action) {
    clearTimer();
    const nextState = actionStates[action] || "idle";
    oneShotState = nextState;
    returnState = "idle";
    setState(nextState);
    const frames = groups[nextState] || groups.idle || [];
    timer = window.setTimeout(finishTimedAction, Math.max(900, frames.length * 130));
  }

  function tick() {
    const frames = groups[state] || groups.idle || [];
    const frame = frames[frameIndex % Math.max(frames.length, 1)] || { row: 0, column: 0 };
    sprite.style.backgroundPosition = `${-Math.round(frame.column * cellWidth * scale)}px ${-Math.round(frame.row * cellHeight * scale)}px`;
    frameIndex += 1;
  }

  function hideMenu() {
    menu.classList.remove("open");
  }

  function showPanel(name) {
    for (const [panelName, panel] of Object.entries(panels)) {
      panel.hidden = panelName !== name;
    }
  }

  function showMenu(clientX, clientY) {
    showPanel("main");
    menu.classList.add("open");
    menu.style.left = "0px";
    menu.style.top = "0px";
    const rootRect = root.getBoundingClientRect();
    const menuRect = menu.getBoundingClientRect();
    const x = clamp(
      root.offsetWidth - menuRect.width + 6,
      -menuRect.width + 28,
      window.innerWidth - rootRect.left - menuRect.width - 8
    );
    const y = clamp(
      clientY - rootRect.top,
      -menuRect.height + 28,
      window.innerHeight - rootRect.top - menuRect.height - 8
    );
    menu.style.left = `${x}px`;
    menu.style.top = `${y}px`;
  }

  function handlePointerDown(event) {
    if (event.target.closest("#acta-web-pet-menu")) return;
    if (event.target.closest("#acta-web-pet-qa")) return;
    pointerDown = true;
    dragging = false;
    pressButton = event.button;
    const rect = root.getBoundingClientRect();
    const frameRect = frame ? frame.getBoundingClientRect() : null;
    offsetX = event.clientX - rect.left;
    offsetY = event.clientY - rect.top;
    framePointerOffsetX = frameRect ? event.clientX : offsetX;
    framePointerOffsetY = frameRect ? event.clientY : offsetY;
    pressX = event.clientX;
    pressY = event.clientY;
    lastPointerX = event.clientX;
    lastScreenX = event.screenX;
    lastScreenY = event.screenY;
    dragState = "running-right";
    hideMenu();
    root.setPointerCapture?.(event.pointerId);
  }

  function handlePointerMove(event) {
    if (!pointerDown) return;
    const distance = Math.abs(event.clientX - pressX) + Math.abs(event.clientY - pressY);
    if (!dragging && distance < 6) return;
    if (clickTimer) {
      window.clearTimeout(clickTimer);
      clickTimer = null;
    }
    if (oneShotState) {
      clearTimer();
      returnState = "idle";
    }
    dragging = true;
    root.classList.add("dragging");
    const deltaX = event.screenX - lastScreenX;
    const deltaY = event.screenY - lastScreenY;
    lastScreenX = event.screenX;
    lastScreenY = event.screenY;
    if (frame) {
      const frameRect = frame.getBoundingClientRect();
      const viewport = hostViewport();
      frame.style.right = "auto";
      frame.style.bottom = "auto";
      frame.style.left = `${clamp(frameRect.left + event.clientX - framePointerOffsetX, 8, viewport.width - frame.offsetWidth - 8)}px`;
      frame.style.top = `${clamp(frameRect.top + event.clientY - framePointerOffsetY, 8, viewport.height - frame.offsetHeight - 8)}px`;
    } else {
      root.style.left = `${clamp(event.clientX - offsetX, 8, window.innerWidth - root.offsetWidth - 8)}px`;
      root.style.top = `${clamp(event.clientY - offsetY, 8, window.innerHeight - root.offsetHeight - 8)}px`;
    }
    lastPointerX = event.clientX;
    if (deltaX < -1) dragState = "running-left";
    if (deltaX > 1) dragState = "running-right";
    setState(dragState);
  }

  function handlePointerUp(event) {
    if (!pointerDown) return;
    const wasDragging = dragging;
    const moved = Math.abs(event.clientX - pressX) + Math.abs(event.clientY - pressY);
    pointerDown = false;
    dragging = false;
    root.classList.remove("dragging");
    root.releasePointerCapture?.(event.pointerId);

    if (wasDragging) {
      setState(returnState);
      return;
    }
    if (pressButton === 2 || event.button === 2) {
      if (moved < 6) showMenu(event.clientX, event.clientY);
      return;
    }
    if (clickTimer) {
      window.clearTimeout(clickTimer);
      clickTimer = null;
      playOnce("jump");
      window.open(config.actaUrl, "_blank", "noopener,noreferrer");
      return;
    }
    clickTimer = window.setTimeout(() => {
      clickTimer = null;
      playOnce("think");
    }, 360);
  }

  root.addEventListener("pointerdown", handlePointerDown);
  root.addEventListener("pointermove", handlePointerMove);
  root.addEventListener("pointerup", handlePointerUp);
  root.addEventListener("pointercancel", handlePointerUp);
  root.addEventListener("contextmenu", (event) => {
    event.preventDefault();
    if (!dragging) showMenu(event.clientX, event.clientY);
  });
  doc.addEventListener("pointerdown", (event) => {
    if (!root.contains(event.target)) hideMenu();
  });
  try {
    window.parent.document.addEventListener("pointerdown", (event) => {
      if (frame && event.target !== frame) hideMenu();
    });
  } catch (error) {
    // Cross-origin embedding would simply keep the in-frame listener above.
  }
  window.addEventListener("resize", () => {
    const rect = root.getBoundingClientRect();
    root.style.left = `${clamp(rect.left, 8, window.innerWidth - root.offsetWidth - 8)}px`;
    root.style.top = `${clamp(rect.top, 8, window.innerHeight - root.offsetHeight - 8)}px`;
  });

  qaSend.addEventListener("click", sendQa);
  qaClose.addEventListener("click", closeQa);
  qaInput.addEventListener("keydown", (event) => {
    if (event.key === "Enter") sendQa();
    if (event.key === "Escape") closeQa();
  });

  menu.addEventListener("click", (event) => {
    const button = event.target.closest("button");
    if (!button) return;
    if (button.dataset.menuTarget) {
      showPanel(button.dataset.menuTarget);
      return;
    }
    if (button.dataset.close) {
      if (talkTimer) window.clearInterval(talkTimer);
      if (speechTimer) window.clearTimeout(speechTimer);
      root.remove();
      style.remove();
      if (frame) frame.style.display = "none";
      return;
    }
    if (button.dataset.talkToggle) {
      talkMode = !talkMode;
      localStorage.setItem("actaPetTalkMode", String(talkMode));
      if (!talkMode) closeQa();
      updateTalkToggleButton();
      refreshTalkTimer(talkMode);
      hideMenu();
      return;
    }
    if (button.dataset.qaMode) {
      openQa();
      hideMenu();
      return;
    }
    if (button.dataset.timedAction) {
      selectedTimedAction = button.dataset.timedAction;
      const labels = {
        rest: "&#20241;&#24687;",
        paper: "&#35770;&#25991;",
        experiment: "&#23454;&#39564;",
      };
      durationTitle.innerHTML = `${labels[selectedTimedAction] || ""} - &#36873;&#25321;&#26102;&#38271;`;
      showPanel("duration");
      return;
    }
    if (button.dataset.duration) {
      if (selectedTimedAction) {
        setTimedAction(selectedTimedAction, Number(button.dataset.duration));
      }
      selectedTimedAction = null;
      hideMenu();
      return;
    }
    if (button.dataset.size) {
      applyScale(button.dataset.size);
      hideMenu();
      return;
    }
    const action = button.dataset.action;
    if (!action) return;
    if (action === "wave") {
      playOnce(action);
    } else {
      setFreeAction(action);
    }
    hideMenu();
  });

  applyScale(scale);
  updateTalkToggleButton();
  setState("idle");
  tick();
  window.setInterval(tick, 120);
  playOnce("wave");
  refreshTalkTimer(false);
})();
</script>
"""
    components.html(
        html.replace("__PET_CONFIG__", json.dumps(pet_config)).replace(
            "<script>", "<!-- acta-web-pet-v15 -->\n<script>"
        ),
        height=1,
    )

# Default prompt templates (auto-created as .docx on first run)
DEFAULT_PROMPT_PAGE1 = (
    "Based on the keywords [{keywords}], generate 5 forward-looking "
    "scientific research questions. "
    "For each question, provide:\n"
    "1. The question itself (beginning with 'How', 'In what ways', "
    "or 'To what extent')\n"
    "2. A set of extended keywords derived from the user's input, "
    "broadening the research scope\n\n"
    "Format each question as:\n"
    "---QUESTION---\n"
    "Q: [question text]\n"
    "EXTENDED: [comma-separated extended keywords]\n"
    "---END---"
)

DEFAULT_PROMPT_PAGE2 = (
    "Generate a rigorous academic paper outline for the paper titled "
    "\"{title}\".\n"
    "Paper type: [{paper_type}]\n"
    "Additional description: {description}\n\n"
    "Requirements:\n"
    "- The outline must contain at most 6 major chapters (level-1 headings).\n"
    "- Each chapter contains level-2 sections.\n"
    "- Must include: Introduction, Related Work, "
    "Methodology/Experimental Design, Results & Analysis, Discussion, Conclusion.\n"
    "- Use Markdown format: # for chapters, ## for sections, "
    "### for subsections.\n"
    "- Keep headings concise and academically rigorous."
)

SYSTEM_PROMPT = (
    "You are a senior research scientist and academic writing expert. "
    "Provide rigorous, insightful, and well-structured academic output. "
    "Always follow the requested format precisely."
)

DEFAULT_PROMPT_PAGE3 = (
    "You are helping a researcher draft the Introduction section of an academic paper. "
    "Based on the topic description below, generate key argument sentences "
    "(thesis statements) for each of the four standard Introduction paragraphs.\n\n"
    "Topic: {topic}\n\n"
    "Structure your output into exactly four labeled sections:\n\n"
    "**Paragraph 1 — Broad Background & Scientific Problem:**\n"
    "Provide 2–3 key thesis sentences that establish the research domain, "
    "its importance, and the core scientific gap. "
    "Where a claim would normally be backed by a reference, insert a placeholder "
    "in exactly this format: [可插入XX类型文献以强化“xxx”论证] "
    "where 'XX类型文献' describes the reference type needed "
    "(e.g., 综述/原创研究/方法论文/基准数据集) "
    "and 'xxx' is a 3–8 word summary of the specific claim to support.\n\n"
    "**Paragraph 2 — Existing Research:**\n"
    "Provide 3–4 key thesis sentences summarizing the landscape of existing work. "
    "Use citation placeholders as described above at every point where "
    "a reference is expected.\n\n"
    "**Paragraph 3 — Limitations of Existing Techniques:**\n"
    "Provide 2–3 key thesis sentences critically analyzing gaps and shortcomings "
    "in prior work. Use citation placeholders where specific limitations "
    "are attributed to existing studies.\n\n"
    "**Paragraph 4 — Our Contribution:**\n"
    "Provide 2–3 key thesis sentences stating the novel approach and "
    "contributions of this paper. Use placeholders only where absolutely necessary "
    "(e.g., referencing baseline methods for comparison).\n\n"
    "CRITICAL RULES:\n"
    "- Output ONLY thesis/argument sentences — do NOT write full paragraphs.\n"
    "- Every citation-requiring claim MUST use the placeholder format exactly:\n"
    "  [可插入XX类型文献以强化“xxx”论证]\n"
    "- Fill 'XX类型文献' and 'xxx' with context-appropriate content.\n"
    "- Do NOT invent any real or fake paper titles, author names, or DOIs.\n"
    "- Do NOT perform any literature search.\n"
    "- Use formal academic English throughout.\n"
    "- Output only the content — no meta-commentary."
)


DEFAULT_PROMPT_PAGE4 = (
    "# About ACTA\n\n"
    "ACTA (ACademic Trash fActory) is an AI-powered research writing assistant "
    "designed to help researchers draft and refine academic papers.\n\n"
    "## Features\n\n"
    "- **Scientific Question Synthesis** — Generate forward-looking research questions from keywords\n"
    "- **Framework Architect** — Build structured paper outlines with hierarchical chapters\n"
    "- **Introduction Generator** — Draft key thesis sentences for Introduction sections with citation placeholders\n\n"
    "## How to Use\n\n"
    "1. Configure your API key and model in the sidebar\n"
    "2. Navigate to the desired tool page\n"
    "3. Enter your research topic or keywords\n"
    "4. Click **Generate** to get AI-assisted output\n\n"
    "## Prompt Customization\n\n"
    "Each tool page is backed by a Word document (`.docx`) in the project folder. "
    "Edit these files to customize the prompts and fine-tune the output format.\n\n"
    "## Contact\n\n"
    "For questions or feedback, please reach out to the project maintainer."
)
DEFAULT_PROMPT_PAGE5 = (
    "You are an academic publishing advisor. Based on the abstract below, "
    "write a brief analysis (approximately 100 words, in English) covering:\n"
    "1. What broad research field this work belongs to (1-2 disciplines).\n"
    "2. Which research_major_direction might be relevant.\n"
    "3. One practical tip for finding suitable journals in our database, "
    "mentioning specific categories from the available columns: "
    "Category (EN) and research_major_direction (specific research area).\n\n"
    "Keep the tone helpful and concise. Do NOT list specific journal names.\n\n"
    "Abstract: {abstract}"
)











# ==================== Helper Functions ====================

def ensure_prompt_doc_exists(filepath, default_content):
    """Create the prompt Word document if it does not already exist."""
    if not os.path.exists(filepath):
        try:
            doc = Document()
            for para_text in default_content.split("\n"):
                doc.add_paragraph(para_text)
            doc.save(filepath)
        except Exception as e:
            st.warning(
                f"Could not auto-create prompt document "
                f"{os.path.basename(filepath)}: {e}"
            )


def load_prompt_from_docx(filepath, fallback):
    """Load the prompt template from a Word document.

    If the document is missing or empty, it is created from the fallback
    string and then re-read.
    """
    ensure_prompt_doc_exists(filepath, fallback)
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            ensure_prompt_doc_exists(filepath, fallback)
            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        st.error(f"Failed to read prompt document: {e}")
        return fallback


def _git_commit_token_update():
    """Best-effort git add + commit + push of user_keys.xlsx after token update.

    Writes status to git_sync.log in the project directory for debugging.
    """
    import datetime
    log_path = os.path.join(APP_DIR, "git_sync.log")

    def _log(msg):
        try:
            with open(log_path, "a", encoding="utf-8") as lf:
                lf.write(f"[{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {msg}\n")
        except Exception:
            pass

    try:
        excel_path = os.path.join(APP_DIR, "user_keys.xlsx")
        repo_dir = APP_DIR

        # Step 1: check git repo
        r = subprocess.run(
            ["git", "rev-parse", "--git-dir"],
            cwd=repo_dir, capture_output=True, text=True, timeout=5
        )
        if r.returncode != 0:
            _log("SKIP: not a git repository")
            return

        # Step 2: git add
        r = subprocess.run(
            ["git", "add", "user_keys.xlsx"],
            cwd=repo_dir, capture_output=True, text=True, timeout=10
        )
        if r.returncode != 0:
            _log(f"FAIL git add: {r.stderr.strip()}")
            return

        # Step 3: check if there are staged changes
        r = subprocess.run(
            ["git", "diff", "--cached", "--quiet", "user_keys.xlsx"],
            cwd=repo_dir, capture_output=True, timeout=10
        )
        if r.returncode == 0:
            _log("SKIP: no changes to commit")
            return

        # Step 4: commit
        msg = f"Update token usage - {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        r = subprocess.run(
            ["git", "commit", "-m", msg],
            cwd=repo_dir, capture_output=True, text=True, timeout=10
        )
        if r.returncode != 0:
            _log(f"FAIL git commit: {r.stderr.strip()}")
            return
        _log(f"COMMIT: {msg}")

        # Step 5: push
        r = subprocess.run(
            ["git", "push"],
            cwd=repo_dir, capture_output=True, text=True, timeout=30
        )
        if r.returncode != 0:
            _log(f"FAIL git push (rc={r.returncode}): {r.stderr.strip()[:200]}")
            _log("HINT: Run 'git push' manually or configure credential.helper")
            return
        _log("PUSH: success")

    except subprocess.TimeoutExpired:
        _log("FAIL: git operation timed out")
    except FileNotFoundError:
        _log("FAIL: git command not found (is Git installed?)")
    except Exception as e:
        _log(f"FAIL: unexpected error - {e}")




def load_user_key_credentials(user_key_input):
    """Look up a 10-digit user key in the password-protected Excel.

    Columns: user_key | api_key | base_url | model | token_limit | token_used
    Returns (api_key, base_url, model, token_limit, token_used, row_index)
    or (None,)*7 if not found.
    """
    excel_path = os.path.join(APP_DIR, "user_keys.xlsx")
    if not os.path.exists(excel_path):
        return None, None, None, None, None, None

    user_key_input = str(user_key_input).strip()
    if len(user_key_input) != 10 or not user_key_input.isdigit():
        return None, None, None, None, None, None

    EXCEL_PASSWORD = "980120"

    try:
        decrypted = io.BytesIO()
        with open(excel_path, "rb") as f:
            office_file = msoffcrypto.OfficeFile(f)
            office_file.load_key(password=EXCEL_PASSWORD)
            office_file.decrypt(decrypted)
        decrypted.seek(0)

        wb = load_workbook(decrypted, read_only=True)
        ws = wb.active
        row_idx = 2  # 1-based, skip header
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] is None:
                row_idx += 1
                continue
            stored_key = str(row[0]).strip()
            if stored_key == user_key_input:
                api_key = str(row[1]).strip() if row[1] else None
                base_url = str(row[2]).strip() if row[2] else None
                model = str(row[3]).strip() if len(row) > 3 and row[3] else None
                token_limit = int(row[4]) if len(row) > 4 and row[4] is not None else 200000
                token_used = int(row[5]) if len(row) > 5 and row[5] is not None else 0
                wb.close()
                return api_key, base_url, model, token_limit, token_used, row_idx
            row_idx += 1
        wb.close()
    except Exception:
        pass
    return None, None, None, None, None, None


def update_token_usage(user_key_input, row_index, new_token_used):
    """Write the updated token_used back to the encrypted Excel for a given key row."""
    excel_path = os.path.join(APP_DIR, "user_keys.xlsx")
    EXCEL_PASSWORD = "980120"

    try:
        # 1. Decrypt
        decrypted = io.BytesIO()
        with open(excel_path, "rb") as f:
            office_file = msoffcrypto.OfficeFile(f)
            office_file.load_key(password=EXCEL_PASSWORD)
            office_file.decrypt(decrypted)
        decrypted.seek(0)

        # 2. Load (not read_only) and update
        wb = load_workbook(decrypted)
        ws = wb.active
        ws.cell(row=row_index, column=6, value=new_token_used)  # column 6 = token_used

        # 3. Save to BytesIO
        saved = io.BytesIO()
        wb.save(saved)
        wb.close()
        saved.seek(0)

        # 4. Re-encrypt and write to disk
        with open(excel_path, "wb") as f:
            encrypted_out = io.BytesIO()
            of_enc = OOXMLFile(saved)
            of_enc.encrypt(EXCEL_PASSWORD, encrypted_out)
            encrypted_out.seek(0)
            f.write(encrypted_out.read())
    except Exception as e:
        # Log write-back failures
        try:
            log_path = os.path.join(APP_DIR, "token_sync.log")
            with open(log_path, "a", encoding="utf-8") as lf:
                import datetime
                lf.write(f"[{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] FAIL update_token_usage: {e}\n")
        except Exception:
            pass


def call_llm(prompt, system_prompt=SYSTEM_PROMPT):
    """Invoke the LLM API with quota check and token tracking.

    Reads API credentials from st.session_state. Supports User Key override
    with token quota enforcement. Raises ValueError when configuration is
    missing or quota is exceeded.
    """
    # API Key takes priority; User Key is fallback
    api_key = st.session_state.get("api_key", "").strip()
    base_url = st.session_state.get("base_url", "").strip()
    model = st.session_state.get("model", "gpt-3.5-turbo").strip()
    token_limit = None
    token_used = None
    row_index = None
    using_user_key = False

    # If API Key is not set, try User Key
    if not api_key:
        user_key = st.session_state.get("user_key", "").strip()
        if user_key:
            uk_api_key, uk_base_url, uk_model, token_limit, token_used, row_index = (
                load_user_key_credentials(user_key)
            )
            if uk_api_key:
                api_key = uk_api_key
                base_url = uk_base_url or "https://api.siliconflow.cn/v1"
                model = uk_model or "deepseek-ai/DeepSeek-V4-Pro"
                using_user_key = True

                # Quota check
                if token_limit is not None and token_used is not None:
                    if token_used >= token_limit:
                        raise ValueError(
                            f"Token quota exhausted ({token_used:,}/{token_limit:,} tokens used). "
                            "Please contact the administrator to increase your quota."
                        )
            else:
                raise ValueError(
                    "Invalid User Key. Please check your 10-digit key or contact the administrator."
                )

    if not api_key:
        raise ValueError("Please configure your API Key or User Key in the sidebar.")
    if not base_url:
        raise ValueError("Please configure the API Base URL in the sidebar.")
    if not model:
        raise ValueError("Please configure the Model Name in the sidebar.")

    # Create HTTP client with timeout
    import httpx
    http_client = httpx.Client(
        timeout=httpx.Timeout(120.0, connect=60.0),
    )

    client = OpenAI(
        api_key=api_key,
        base_url=base_url,
        http_client=http_client,
    )

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ],
        temperature=0.7,
        max_tokens=4096,
        timeout=httpx.Timeout(120.0, connect=60.0),
    )

    # Update token usage for user-key users
    if using_user_key and row_index is not None and token_used is not None:
        try:
            usage = response.usage
            if usage and usage.total_tokens:
                new_total = token_used + usage.total_tokens
                update_token_usage(user_key, row_index, new_total)
                # Update sidebar display
                st.session_state["user_token_used"] = new_total
                # Git auto-commit (best-effort, non-blocking)
                _git_commit_token_update()
        except Exception:
            pass

    return response.choices[0].message.content



def parse_questions(raw_text):
    """Parse LLM output into structured question entries.

    Expected format per question:
        ---QUESTION---
        Q: <question text>
        KEYWORDS: <comma-separated>
        EXTENDED: <comma-separated>
        ---END---

    Returns a list of dicts with keys: question, keywords, extended.
    """
    entries = []
    blocks = re.split(r"---QUESTION---", raw_text)
    for block in blocks:
        block = block.strip()
        if not block:
            continue

        q_match = re.search(r"Q:\s*(.+?)(?:\n|---END)", block, re.DOTALL)
        kw_match = re.search(
            r"KEYWORDS:\s*(.+?)(?:\n|---END|EXTENDED)",
            block, re.IGNORECASE | re.DOTALL,
        )
        ex_match = re.search(
            r"EXTENDED:\s*(.+?)(?:\n|---END|$)",
            block, re.IGNORECASE | re.DOTALL,
        )

        if q_match:
            question = q_match.group(1).strip()
            keywords_raw = kw_match.group(1).strip() if kw_match else ""
            extended_raw = ex_match.group(1).strip() if ex_match else ""

            keywords = [
                k.strip() for k in keywords_raw.split(",") if k.strip()
            ]
            extended = [
                k.strip() for k in extended_raw.split(",") if k.strip()
            ]

            entries.append({
                "question": question,
                "keywords": keywords,
                "extended": extended,
            })

    return entries


def render_framework_markdown(md_text):
    """Render paper framework with indentation and font-size by heading level.

    Heading mapping:
        #   (level 1) — 26px, bold, no indent   (chapter)
        ##  (level 2) — 21px, bold, 20px indent (section)
        ### (level 3) — 17px, semi-bold, 40px indent (subsection)
        ####+       — 14–15px, 56px indent

    Body paragraphs are rendered with comfortable line-height.
    """
    lines = md_text.strip().split("\n")
    html_parts = []

    heading_styles = {
        1: {"size": 26, "weight": 700, "indent": 0,  "color": "#0f172a"},
        2: {"size": 21, "weight": 700, "indent": 20, "color": "#1e293b"},
        3: {"size": 17, "weight": 600, "indent": 40, "color": "#334155"},
        4: {"size": 15, "weight": 600, "indent": 56, "color": "#475569"},
        5: {"size": 14, "weight": 500, "indent": 56, "color": "#64748b"},
        6: {"size": 13, "weight": 500, "indent": 56, "color": "#64748b"},
    }

    for line in lines:
        stripped = line.strip()
        if not stripped:
            html_parts.append('<div style="height:6px;"></div>')
            continue

        # Detect Markdown heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            style = heading_styles.get(level, heading_styles[6])

            html_parts.append(
                f'<div style="'
                f'margin-left:{style["indent"]}px; '
                f'font-size:{style["size"]}px; '
                f'font-weight:{style["weight"]}; '
                f'color:{style["color"]}; '
                f'margin-top:{28 if level <= 2 else 20}px; '
                f'margin-bottom:{10 if level <= 2 else 8}px; '
                f'line-height:1.35; '
                f'letter-spacing:-0.01em;'
                f'">{title}</div>'
            )
            continue

        # Bold standalone paragraph
        bold_match = re.match(r"^\*\*(.+)\*\*$", stripped)
        if bold_match:
            html_parts.append(
                f'<p style="margin-left:20px; font-weight:600; '
                f'color:#334155; line-height:1.7; '
                f'margin-bottom:4px;">{bold_match.group(1)}</p>'
            )
            continue

        # Regular body paragraph
        html_parts.append(
            f'<p style="margin-left:18px; color:#475569; '
            f'line-height:1.8; margin-bottom:6px; '
            f'font-size:15px;">{stripped}</p>'
        )

    return "\n".join(html_parts)


# ==================== Journal Recommender Helpers ====================

JOURNAL_HEADERS = [
    "journal_id", "journal_name", "impact_factor", "citescore",
    "xr_zone", "top", "category_en", "category",
    "research_major_direction", "research_minor_direction",
    "is_oa", "language", "publisher", "self_citation_rate",
    "five_year_if", "h_index", "journal_intro", "website",
    "research_direction", "country", "period", "pub_year",
    "gold_oa_ratio", "research_article_ratio", "review_speed",
    "acceptance_rate", "annual_articles", "citescore_detail",
    "p_issn", "e_issn"
]

JOURNAL_DISPLAY_LABELS = {
    "journal_id": "Journal ID",
    "journal_name": "Journal Name",
    "impact_factor": "Impact Factor",
    "citescore": "CiteScore",
    "xr_zone": "New Sharp Zone",
    "top": "TOP",
    "category_en": "Category (EN)",
    "category": "Category",
    "research_major_direction": "Major Research Direction",
    "research_minor_direction": "Minor Research Direction",
    "is_oa": "Open Access",
    "language": "Language",
    "publisher": "Publisher",
    "self_citation_rate": "Self-citation Rate",
    "five_year_if": "5-Year IF",
    "h_index": "H-index",
    "journal_intro": "Introduction",
    "website": "Website",
    "research_direction": "Research Direction",
    "country": "Country",
    "period": "Period",
    "pub_year": "Publication Year",
    "gold_oa_ratio": "Gold OA Ratio",
    "research_article_ratio": "Research Article Ratio",
    "review_speed": "Review Speed",
    "acceptance_rate": "Acceptance Rate",
    "annual_articles": "Annual Articles",
    "citescore_detail": "CiteScore Detail",
    "p_issn": "P-ISSN",
    "e_issn": "E-ISSN",
}

FILTER_COLUMNS = {"category", "research_major_direction", "impact_factor", "citescore", "xr_zone", "top"}
HEADER_COLUMNS = {"journal_name", "impact_factor", "xr_zone"}
DISPLAY_COLUMNS = [h for h in JOURNAL_HEADERS if h not in HEADER_COLUMNS]


def load_journal_data():
    jdata_path = os.path.join(APP_DIR, "Jdata_with_cn_v1.xlsx")
    wb = load_workbook(jdata_path, read_only=True)
    ws = wb.active
    journals = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[8] is not None and str(row[8]).strip() != "":
            d = {h: (row[i] if i < len(row) else None) for i, h in enumerate(JOURNAL_HEADERS)}
            journals.append(d)
    wb.close()
    return journals


def get_unique_values(journals, col):
    """Get sorted unique non-null values for a column."""
    vals = set()
    for j in journals:
        v = j.get(col)
        if v is not None and str(v).strip() != "" and str(v).strip().lower() != "none":
            vals.add(str(v).strip())
    try:
        return sorted(vals, key=lambda x: (
            float(x) if x.replace(".", "").replace("-", "").lstrip("-").isdigit() else float("inf"), x
        ))
    except Exception:
        return sorted(vals)


def filter_journals_by_selection(journals, selections, skip_key=None):
    """Filter journals by selected values, optionally skipping one filter key."""
    filtered = []
    for journal in journals:
        keep = True
        for key, selected_values in selections.items():
            if key == skip_key or not selected_values:
                continue
            journal_value = str(journal.get(key) or "").strip()
            if journal_value not in selected_values:
                keep = False
                break
        if keep:
            filtered.append(journal)
    return filtered


def normalize_filter_values(raw_value):
    """Normalize a filter value into a clean list of strings."""
    if isinstance(raw_value, str):
        raw_value = [raw_value] if raw_value else []
    elif raw_value is None:
        raw_value = []
    return [str(v).strip() for v in raw_value if str(v).strip()]


def build_multiselect_options(options, selected_values):
    """Keep selected values visible while showing current available options."""
    merged = list(selected_values)
    for option in options:
        if option not in merged:
            merged.append(option)
    return merged


def find_journal_matches(journals, query, limit=5):
    """Return up to `limit` closest journal matches for a free-text query."""
    query_norm = str(query or "").strip().lower()
    if not query_norm:
        return []

    exact_matches = []
    ranked_matches = []
    for journal in journals:
        name = str(journal.get("journal_name") or "").strip()
        if not name:
            continue
        name_norm = name.lower()
        if name_norm == query_norm:
            exact_matches.append(journal)
            continue
        score = difflib.SequenceMatcher(None, query_norm, name_norm).ratio()
        if query_norm in name_norm:
            score += 0.35
        if name_norm.startswith(query_norm):
            score += 0.25
        if score > 0.2:
            ranked_matches.append((score, name, journal))

    if exact_matches:
        return exact_matches[:1]

    ranked_matches.sort(key=lambda item: (-item[0], item[1].lower()))
    return [journal for _, _, journal in ranked_matches[:limit]]


PAGE5_ALLOWED_FIELDS = [
    "Medicine", "Social Sciences", "Engineering", "Literature and Language",
    "Biology", "Management Science", "Agricultural and Forestry Science",
    "Mathematics", "Education Science", "Philosophy", "Economics", "History",
    "Computer Science", "Psychology", "GeoSciences", "Chemistry",
    "Environment Science and Ecology", "Materials Science", "Art",
    "Physics and Astronomy", "Multidisciplinary Science",
]

PAGE5_FIELD_EXPANSIONS = {
    "Engineering": [
        "engineering", "civil engineering", "geotechnical engineering",
        "geological engineering", "underground engineering", "tunnel",
        "infrastructure", "site investigation", "soil mechanics",
        "rock mechanics", "borehole", "foundation", "slope",
    ],
    "GeoSciences": [
        "geosciences", "geology", "engineering geology", "geotechnical",
        "stratigraphy", "stratigraphic", "hydrogeology", "geomorphology",
        "landslide", "geological hazard", "borehole", "subsurface",
        "sediment", "rock", "soil", "earth science",
    ],
    "Environment Science and Ecology": [
        "environmental science", "ecology", "sustainability", "climate",
        "pollution", "environmental management", "natural hazard",
        "risk assessment",
    ],
    "Computer Science": [
        "computer science", "artificial intelligence", "machine learning",
        "deep learning", "data mining", "algorithm", "computer vision",
        "pattern recognition", "informatics",
    ],
    "Materials Science": [
        "materials science", "materials", "composite", "nanomaterial",
        "polymer", "ceramic", "metallurgy",
    ],
    "Medicine": [
        "medicine", "clinical", "public health", "biomedical", "patient",
        "disease", "therapy", "diagnosis",
    ],
}

PAGE5_TOPIC_RULES = [
    {
        "abstract_terms": [
            "stratigraph", "borehole", "geotechnical", "engineering geology",
            "geological site", "underground", "tunnel", "soil", "rock",
        ],
        "fields": ["Engineering", "GeoSciences"],
        "majors": ["GeoSciences", "Engineering", "Geology", "Geotechnical"],
        "directions": [
            "engineering geology", "geotechnical engineering",
            "stratigraphic uncertainty", "borehole sampling",
            "underground design", "site investigation", "geological modelling",
        ],
    },
    {
        "abstract_terms": ["landslide", "slope failure", "debris flow"],
        "fields": ["GeoSciences", "Engineering", "Environment Science and Ecology"],
        "majors": ["GeoSciences", "Engineering", "Environmental Science"],
        "directions": [
            "landslide", "slope stability", "geological hazard",
            "engineering geology", "natural hazards", "risk assessment",
        ],
    },
    {
        "abstract_terms": ["machine learning", "deep learning", "neural network", "artificial intelligence"],
        "fields": ["Computer Science"],
        "majors": ["Computer", "Informatics", "Statistics"],
        "directions": [
            "artificial intelligence", "machine learning", "data mining",
            "pattern recognition", "statistical modelling",
        ],
    },
]


def _page5_clean_text(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _page5_extract_json(text):
    """Extract the first JSON object from an LLM response."""
    raw = str(text or "").strip()
    if not raw:
        return {}
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    try:
        return json.loads(raw)
    except Exception:
        pass
    match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except Exception:
            return {}
    return {}


def _page5_list(value):
    if isinstance(value, list):
        return [_page5_clean_text(v) for v in value if _page5_clean_text(v)]
    if isinstance(value, str) and value.strip():
        parts = re.split(r"[,;|\n]+", value)
        return [_page5_clean_text(v) for v in parts if _page5_clean_text(v)]
    return []


def _page5_match_allowed_field(value):
    value_norm = _page5_clean_text(value).lower()
    if not value_norm:
        return None
    for field in PAGE5_ALLOWED_FIELDS:
        field_norm = field.lower()
        if field_norm == value_norm or field_norm in value_norm or value_norm in field_norm:
            return field
    if "earth" in value_norm or "geo" in value_norm or "geology" in value_norm:
        return "GeoSciences"
    if "environment" in value_norm or "ecology" in value_norm:
        return "Environment Science and Ecology"
    if "computer" in value_norm or "artificial" in value_norm or "informatics" in value_norm:
        return "Computer Science"
    return None


def _page5_add_unique(target, values):
    for value in values:
        clean = _page5_clean_text(value)
        if clean and clean not in target:
            target.append(clean)


def _page5_apply_topic_rules(analysis, abstract):
    abstract_l = str(abstract or "").lower()
    for rule in PAGE5_TOPIC_RULES:
        if any(term in abstract_l for term in rule["abstract_terms"]):
            _page5_add_unique(analysis["broad_fields"], rule["fields"])
            _page5_add_unique(analysis["major_directions"], rule["majors"])
            _page5_add_unique(analysis["directions"], rule["directions"])
    for field in list(analysis["broad_fields"]):
        _page5_add_unique(analysis["directions"], PAGE5_FIELD_EXPANSIONS.get(field, []))
    return analysis


def _page5_fallback_analysis(abstract):
    analysis = {
        "broad_fields": [],
        "major_directions": [],
        "directions": [],
        "semantic_query": _page5_clean_text(abstract),
        "reasoning": "",
    }
    abstract_l = str(abstract or "").lower()
    for field, terms in PAGE5_FIELD_EXPANSIONS.items():
        if any(term in abstract_l for term in terms):
            _page5_add_unique(analysis["broad_fields"], [field])
    if not analysis["broad_fields"]:
        _page5_add_unique(analysis["broad_fields"], ["Multidisciplinary Science"])
    return _page5_apply_topic_rules(analysis, abstract)


def build_page5_analysis_prompt(abstract):
    fields = ", ".join(PAGE5_ALLOWED_FIELDS)
    return (
        "You are an academic journal recommendation classifier. Read the abstract "
        "and identify the journals' target domain, not merely the method domain. "
        "If machine learning is used to solve a geology, hazard, medical, or "
        "engineering problem, prioritize the application field that receives the "
        "main contribution.\n\n"
        "Return ONLY valid JSON with these keys:\n"
        "{\n"
        '  "broad_fields": ["1-3 items from the allowed list"],\n'
        '  "major_directions": ["3-8 database direction keywords"],\n'
        '  "directions": ["6-12 topical phrases for semantic retrieval"],\n'
        '  "semantic_query": "one expanded retrieval query with synonyms",\n'
        '  "reasoning": "one concise sentence"\n'
        "}\n\n"
        "Allowed broad_fields: " + fields + "\n\n"
        "Important example: work on 3D stratigraphic uncertainty, boreholes, "
        "geotechnical sites, tunnels, slopes, or underground design should include "
        "Engineering and GeoSciences, with directions such as engineering geology "
        "and geotechnical engineering.\n\n"
        "Abstract:\n" + abstract
    )


def analyze_abstract_for_journals(abstract):
    """Use the LLM for field understanding, then normalize for local retrieval."""
    raw = call_llm(build_page5_analysis_prompt(abstract), system_prompt=SYSTEM_PROMPT)
    data = _page5_extract_json(raw)
    if not data:
        data = {}

    analysis = {
        "broad_fields": [],
        "major_directions": _page5_list(data.get("major_directions")),
        "directions": _page5_list(data.get("directions")),
        "semantic_query": _page5_clean_text(data.get("semantic_query")) or _page5_clean_text(abstract),
        "reasoning": _page5_clean_text(data.get("reasoning")) or _page5_clean_text(raw),
        "raw_response": raw,
    }
    for field in _page5_list(data.get("broad_fields")):
        matched = _page5_match_allowed_field(field)
        if matched and matched not in analysis["broad_fields"]:
            analysis["broad_fields"].append(matched)

    if not analysis["broad_fields"]:
        fallback = _page5_fallback_analysis(abstract)
        analysis["broad_fields"] = fallback["broad_fields"]
        _page5_add_unique(analysis["major_directions"], fallback["major_directions"])
        _page5_add_unique(analysis["directions"], fallback["directions"])

    analysis = _page5_apply_topic_rules(analysis, abstract)
    query_parts = [abstract, analysis["semantic_query"]]
    query_parts.extend(analysis["broad_fields"])
    query_parts.extend(analysis["major_directions"])
    query_parts.extend(analysis["directions"])
    analysis["retrieval_query"] = " ".join(_page5_clean_text(p) for p in query_parts if _page5_clean_text(p))
    return analysis


def _page5_zone_key(value):
    try:
        return str(int(float(value)))
    except Exception:
        return _page5_clean_text(value)


def _page5_journal_text(journal):
    weighted_parts = []
    for key, repeat in [
        ("journal_name", 5),
        ("category_en", 3),
        ("research_major_direction", 4),
        ("research_minor_direction", 4),
        ("research_direction", 3),
        ("journal_intro", 3),
        ("category", 1),
    ]:
        value = _page5_clean_text(journal.get(key))
        if value:
            weighted_parts.extend([value] * repeat)
    return " ".join(weighted_parts)


def _page5_field_match(journal, analysis):
    fields = analysis.get("broad_fields", [])
    majors = [m.lower() for m in analysis.get("major_directions", [])]
    directions = [d.lower() for d in analysis.get("directions", [])]
    category = _page5_clean_text(journal.get("category_en")).lower()
    major = _page5_clean_text(journal.get("research_major_direction")).lower()
    minor = _page5_clean_text(journal.get("research_minor_direction")).lower()
    research_direction = _page5_clean_text(journal.get("research_direction")).lower()
    intro = _page5_clean_text(journal.get("journal_intro")).lower()
    combined = " ".join([category, major, minor, research_direction, intro])

    if any(field.lower() == category for field in fields):
        return True
    for field in fields:
        if field.lower() in major or field.lower() in research_direction:
            return True
    for token in majors:
        if token and (token in major or token in minor or token in research_direction):
            return True
    high_value_directions = [d for d in directions if len(d) >= 6]
    if any(d in combined for d in high_value_directions):
        return True
    return False


def _page5_rule_boost(journal, analysis, abstract):
    text = _page5_journal_text(journal).lower()
    name = _page5_clean_text(journal.get("journal_name")).lower()
    query = (str(abstract or "") + " " + analysis.get("retrieval_query", "")).lower()
    boost = 0.0

    if any(t in query for t in ["stratigraph", "borehole", "geotechnical", "underground", "tunnel"]):
        if "engineering geology" in name:
            boost += 0.45
        if "engineering geology" in text:
            boost += 0.25
        if "geotechnical" in text:
            boost += 0.12
    if "landslide" in query or "slope" in query:
        if any(t in text for t in ["landslide", "slope stability", "engineering geology", "geological hazard"]):
            boost += 0.20
    if any(t in query for t in ["machine learning", "deep learning", "artificial intelligence"]):
        if any(t in text for t in ["artificial intelligence", "machine learning", "pattern recognition", "informatics"]):
            boost += 0.12
    for direction in analysis.get("directions", []):
        d = direction.lower()
        if len(d) >= 8 and d in text:
            boost += 0.04
    return min(boost, 0.65)


def recommend_journals_from_abstract(journals, analysis, abstract, per_zone=10):
    """Recommend top journals in XINRUI zones 1/2/3 using semantic retrieval."""
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
    except Exception as e:
        raise RuntimeError("scikit-learn is required for journal recommendation: " + str(e))

    zone_results = {}
    query = analysis.get("retrieval_query") or abstract

    for zone in ["1", "2", "3"]:
        zone_pool = [j for j in journals if _page5_zone_key(j.get("xr_zone")) == zone]
        field_pool = [j for j in zone_pool if _page5_field_match(j, analysis)]
        candidate_pool = field_pool if len(field_pool) >= per_zone else zone_pool
        if not candidate_pool:
            zone_results[zone] = []
            continue

        docs = [_page5_journal_text(j) for j in candidate_pool]
        vectorizer = TfidfVectorizer(
            lowercase=True,
            stop_words="english",
            ngram_range=(1, 2),
            min_df=1,
            sublinear_tf=True,
        )
        matrix = vectorizer.fit_transform([query] + docs)
        semantic_scores = cosine_similarity(matrix[0:1], matrix[1:]).ravel()

        ranked = []
        for idx, journal in enumerate(candidate_pool):
            field_boost = 0.12 if journal in field_pool else 0.0
            rule_boost = _page5_rule_boost(journal, analysis, abstract)
            if_val = journal.get("impact_factor") or 0
            cs_val = journal.get("citescore") or 0
            try:
                prestige = min(math.log1p(float(if_val)) / 10.0, 0.10)
            except Exception:
                prestige = 0.0
            try:
                prestige += min(math.log1p(float(cs_val)) / 20.0, 0.06)
            except Exception:
                pass
            score = float(semantic_scores[idx]) + field_boost + rule_boost + prestige
            item = dict(journal)
            item["recommend_score"] = round(score, 4)
            item["semantic_score"] = round(float(semantic_scores[idx]), 4)
            item["match_reason"] = _page5_build_match_reason(item, analysis, abstract, rule_boost)
            ranked.append(item)

        ranked.sort(key=lambda j: (-j.get("recommend_score", 0), str(j.get("journal_name") or "")))
        zone_results[zone] = ranked[:per_zone]
    return zone_results


def _page5_build_match_reason(journal, analysis, abstract, rule_boost):
    name = _page5_clean_text(journal.get("journal_name"))
    category = _page5_clean_text(journal.get("category_en"))
    major = _page5_clean_text(journal.get("research_major_direction"))
    parts = []
    if category:
        parts.append(category)
    if major and major.lower() not in category.lower():
        parts.append(major)
    if rule_boost >= 0.25:
        parts.append("strong topical fit")
    elif rule_boost > 0:
        parts.append("topical keyword expansion")
    reason = " / ".join(parts[:3])
    return reason or name or "semantic match"


def render_page5_recommendations(recommendations):
    zone_labels = {"1": "Zone 1", "2": "Zone 2", "3": "Zone 3"}
    for zone in ["1", "2", "3"]:
        rows = recommendations.get(zone, [])
        st.markdown("#### " + zone_labels[zone] + " Recommendations")
        if not rows:
            st.info("No matching journals found in this zone.")
            continue
        for rank, journal in enumerate(rows, 1):
            prefix = str(rank) + ". "
            score = journal.get("recommend_score")
            name = _page5_clean_text(journal.get("journal_name")) or "Unknown"
            label = prefix + name
            if score is not None:
                label += " | relevance " + str(score)
            with st.expander(label, expanded=(rank <= 3)):
                c1, c2, c3 = st.columns(3)
                with c1:
                    st.markdown("**Impact Factor:** " + str(journal.get("impact_factor") or "N/A"))
                with c2:
                    st.markdown("**CiteScore:** " + str(journal.get("citescore") or "N/A"))
                with c3:
                    st.markdown("**New Sharp Zone:** " + str(journal.get("xr_zone") or "N/A"))
                st.markdown("**Match reason:** " + _page5_clean_text(journal.get("match_reason")))
                for key in ["category_en", "research_major_direction", "research_minor_direction", "journal_intro", "website"]:
                    value = journal.get(key)
                    if value is None or str(value).strip() == "":
                        continue
                    label_name = JOURNAL_DISPLAY_LABELS.get(key, key)
                    if key == "website":
                        st.markdown("**" + label_name + ":** [" + str(value) + "](" + str(value) + ")")
                    else:
                        st.markdown("**" + label_name + ":** " + str(value))


def render_journal_entries(journals):
    """Render journal rows as expandable entries."""
    for j in journals:
        name = str(j.get("journal_name") or "Unknown")
        impact = j.get("impact_factor")
        if_str = ("%.2f" % impact) if impact is not None else "N/A"
        xr = j.get("xr_zone")
        xr_str = (" | 新锐分区 " + str(xr)) if xr is not None else ""
        label = name + " (IF: " + if_str + ")" + xr_str
        with st.expander(label):
            for col_key in DISPLAY_COLUMNS:
                val = j.get(col_key)
                if val is not None and str(val).strip() != "" and str(val).strip().lower() != "none":
                    lbl = JOURNAL_DISPLAY_LABELS.get(col_key, col_key)
                    if col_key == "website" and val:
                        st.markdown("**" + lbl + ":** [" + str(val) + "](" + str(val) + ")")
                    else:
                        st.markdown("**" + lbl + ":** " + str(val))


def prune_journal_filter_chain(journals, filter_specs, selections):
    """Prune staged filter selections and compute downstream options."""
    pruned = {}
    options_map = {}
    current_pool = journals
    for data_key, _, _ in filter_specs:
        options = get_unique_values(current_pool, data_key)
        valid_selected = [value for value in selections.get(data_key, []) if value in options]
        pruned[data_key] = valid_selected
        options_map[data_key] = build_multiselect_options(options, valid_selected)
        if valid_selected:
            current_pool = filter_journals_by_selection(current_pool, {data_key: valid_selected})
    return pruned, options_map


def apply_journal_filter_stage(stage_key):
    """Commit one filter stage and prune only downstream stages."""
    filter_specs = st.session_state.get("page5_filter_specs", [])
    journals_all = st.session_state.get("cached_journals", [])
    if not filter_specs or not journals_all:
        return

    current_filters = st.session_state.get("page5_filters", {})
    selections = {
        data_key: normalize_filter_values(current_filters.get(data_key, []))
        for data_key, _, _ in filter_specs
    }
    widget_key = "page5_" + stage_key + "_widget"
    selections[stage_key] = normalize_filter_values(st.session_state.get(widget_key, []))

    pruned, _ = prune_journal_filter_chain(journals_all, filter_specs, selections)
    st.session_state["page5_filters"] = pruned
    for data_key, _, _ in filter_specs:
        st.session_state["page5_" + data_key + "_widget"] = list(pruned[data_key])
    st.session_state["page5_page"] = 1

def render_journal_results():
    "Render journal list with linked multi-select filters and expandable rows."
    journals_all = st.session_state.get("cached_journals", [])
    if not journals_all:
        st.info("Loading journal database ...")
        with st.spinner("Loading journals ..."):
            st.session_state["cached_journals"] = load_journal_data()
        st.rerun()
    if not journals_all:
        st.info("Journal database not loaded. Please reload the page.")
        return

    filter_specs = [
        ("category", "Category", "page5_cat"),
        ("research_major_direction", "Research Major Direction", "page5_major"),
        ("xr_zone", "新锐分区", "page5_xr"),
        ("top", "TOP", "page5_top"),
    ]
    st.session_state["page5_filter_specs"] = filter_specs
    if "page5_filters" not in st.session_state:
        st.session_state["page5_filters"] = {data_key: [] for data_key, _, _ in filter_specs}

    stored_filters = st.session_state["page5_filters"]
    selections = {data_key: normalize_filter_values(stored_filters.get(data_key, [])) for data_key, _, _ in filter_specs}

    title_col, mode_col = st.columns([2.3, 1.7])
    with title_col:
        st.markdown("### Filter Journals")
    with mode_col:
        search_mode = st.radio(
            "Search Mode",
            ["Category Search", "Journal Search"],
            key="page5_search_mode",
            horizontal=True,
            label_visibility="collapsed",
        )

    if search_mode == "Journal Search":
        journal_query = st.text_input(
            "Journal Name",
            key="page5_journal_query",
            placeholder="Type a journal name",
        )
        if not str(journal_query or "").strip():
            st.info("Enter a journal name to search for the closest match.")
            return

        journal_matches = find_journal_matches(journals_all, journal_query, limit=5)
        if not journal_matches:
            st.info("No journal match found.")
            return

        st.markdown("**Showing " + str(len(journal_matches)) + " journal match(es)**")
        render_journal_entries(journal_matches)
        return

    applied_selections, filter_options = prune_journal_filter_chain(journals_all, filter_specs, selections)
    if applied_selections != selections:
        st.session_state["page5_filters"] = applied_selections
        selections = applied_selections
    else:
        selections = applied_selections

    for data_key, _, _ in filter_specs:
        widget_key = "page5_" + data_key + "_widget"
        if widget_key not in st.session_state:
            st.session_state[widget_key] = list(selections[data_key])
        elif normalize_filter_values(st.session_state.get(widget_key, [])) != selections[data_key]:
            st.session_state[widget_key] = list(selections[data_key])

    c1, c2 = st.columns(2)
    with c1:
        sel_cat = st.multiselect(
            "Category",
            filter_options["category"],
            key="page5_category_widget",
            on_change=apply_journal_filter_stage,
            args=("category",),
        )
    with c2:
        sel_major = st.multiselect(
            "Research Major Direction",
            filter_options["research_major_direction"],
            key="page5_research_major_direction_widget",
            on_change=apply_journal_filter_stage,
            args=("research_major_direction",),
        )

    c3, c4, c5 = st.columns(3)
    with c3:
        sel_xr = st.multiselect(
            "新锐分区",
            filter_options["xr_zone"],
            key="page5_xr_zone_widget",
            on_change=apply_journal_filter_stage,
            args=("xr_zone",),
        )
    with c4:
        sel_top = st.multiselect(
            "TOP",
            filter_options["top"],
            key="page5_top_widget",
            on_change=apply_journal_filter_stage,
            args=("top",),
        )
    with c5:
        sort_options = ["IF: High to Low", "IF: Low to High", "CiteScore: High to Low", "CiteScore: Low to High"]
        sel_sort = st.selectbox("Sort By", sort_options, key="page5_sort")

    sel_cat = selections["category"]
    sel_major = selections["research_major_direction"]
    sel_xr = selections["xr_zone"]
    sel_top = selections["top"]

    current_signature = (
        tuple(sorted(sel_cat)),
        tuple(sorted(sel_major)),
        tuple(sorted(sel_xr)),
        tuple(sorted(sel_top)),
    )
    if st.session_state.get("page5_filter_signature") != current_signature:
        st.session_state["page5_filter_signature"] = current_signature
        st.session_state["page5_page"] = 1

    filtered = filter_journals_by_selection(
        journals_all,
        {
            "category": sel_cat,
            "research_major_direction": sel_major,
            "xr_zone": sel_xr,
            "top": sel_top,
        },
    )

    if "CiteScore" in sel_sort:
        if "High to Low" in sel_sort:
            filtered.sort(key=lambda x: -(x.get("citescore") or 0))
        else:
            filtered.sort(key=lambda x: (x.get("citescore") or 999999))
    else:
        if "High to Low" in sel_sort:
            filtered.sort(key=lambda x: -(x.get("impact_factor") or 0))
        else:
            filtered.sort(key=lambda x: (x.get("impact_factor") or 999999))

    PAGE_SIZE = 20
    total_pages = max(1, (len(filtered) + PAGE_SIZE - 1) // PAGE_SIZE)
    if "page5_page" not in st.session_state:
        st.session_state["page5_page"] = 1
    pg = st.session_state["page5_page"]
    if pg > total_pages:
        pg = total_pages
        st.session_state["page5_page"] = pg
    start_idx = (pg - 1) * PAGE_SIZE
    end_idx = min(start_idx + PAGE_SIZE, len(filtered))

    st.markdown("**Showing " + str(len(filtered)) + " journal(s) (Page " + str(pg) + " of " + str(total_pages) + ")**")

    if not filtered:
        st.info("No journals match the current filters. Try adjusting your selection.")
        return

    render_journal_entries(filtered[start_idx:end_idx])

    if total_pages > 1:
        pc1, pc2, pc3, pc4, pc5 = st.columns([1, 2, 1, 2, 1])
        with pc2:
            if st.button("< Previous", disabled=(pg <= 1), key="pg5_prev", use_container_width=True):
                st.session_state["page5_page"] = max(1, pg - 1)
                st.rerun()
        with pc4:
            if st.button("Next >", disabled=(pg >= total_pages), key="pg5_next", use_container_width=True):
                st.session_state["page5_page"] = min(total_pages, pg + 1)
                st.rerun()

# ==================== Sidebar ====================
with st.sidebar:
    # Hero-style branding — prominent site title
    st.markdown(
        '<div style="font-size:50px; font-weight:800; color:#0f172a; '
        'letter-spacing:-0.02em; margin-bottom:0px;">ACTA</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div style="font-size:13px; color:#94a3b8; font-style:italic; '
        'letter-spacing:0.05em; margin-bottom:2px;">ACademic Trash fActory</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div style="font-size:13px; color:#64748b; '
        'margin-bottom:8px;">AI-Powered Research Assistant</div>',
        unsafe_allow_html=True,
    )
    st.markdown("---")

    # Page navigation
    page = st.radio(
        "Navigation",
        [
            "📋 Journal Recommender",
            "🔬 Scientific Question Synthesis",
            "📐 Framework Architect",
            "📝 Introduction Generator",
            "ℹ️ About ACTA",
        ],
        label_visibility="collapsed",
    )

    st.markdown("---")

    # API configuration (collapsible)
    with st.expander("⚙️ API Configuration", expanded=False):
        user_key = st.text_input(
            "User Key",
            type="password",
            value=st.session_state.get("user_key", ""),
            placeholder="10-digit key (if provided)",
            help="Enter a 10-digit user key to use the internal API. Leave blank to use your own API Key below.",
        )
        st.session_state["user_key"] = user_key

        # Validate user key and show token status
        if user_key and len(str(user_key).strip()) == 10 and str(user_key).strip().isdigit():
            uk_api, uk_url, uk_model, tok_limit, tok_used, _ = load_user_key_credentials(user_key)
            if uk_api and tok_limit is not None:
                st.session_state["user_key_valid"] = True
                st.session_state["user_token_limit"] = tok_limit
                st.session_state["user_token_used"] = tok_used
                remaining = max(0, tok_limit - tok_used)
                pct = (tok_used / tok_limit * 100) if tok_limit > 0 else 0
                if pct >= 90:
                    color = "#ef4444"
                elif pct >= 50:
                    color = "#f59e0b"
                else:
                    color = "#22c55e"
                st.markdown(
                    f'<p style="font-size:11px; color:{color}; margin:-8px 0 8px 0;">'
                    f"Remaining: {remaining:,} / {tok_limit:,} tokens ({100-pct:.0f}%)"
                    f"</p>",
                    unsafe_allow_html=True,
                )
            else:
                st.session_state["user_key_valid"] = False
                st.markdown(
                    '<p style="font-size:11px; color:#ef4444; margin:-8px 0 8px 0;">'
                    "Invalid User Key</p>",
                    unsafe_allow_html=True,
                )
        elif user_key:
            st.session_state["user_key_valid"] = False
            st.markdown(
                '<p style="font-size:11px; color:#ef4444; margin:-8px 0 8px 0;">'
                "Key must be 10 digits</p>",
                unsafe_allow_html=True,
            )


        # Check git sync status
        git_log_path = os.path.join(APP_DIR, "git_sync.log")
        git_warn = False
        if os.path.exists(git_log_path):
            try:
                with open(git_log_path, "r", encoding="utf-8") as lf:
                    lines = lf.readlines()
                    if lines:
                        last = lines[-1]
                        if "FAIL" in last:
                            git_warn = True
                            st.warning(
                                "GitHub sync may not be working. "
                                "Check git_sync.log in the project folder."
                            )
            except Exception:
                pass

        st.markdown(
            '<p style="font-size:11px; color:#94a3b8; margin:-10px 0 10px 0;">'
            "— OR —"
            "</p>",
            unsafe_allow_html=True,
        )

        api_key = st.text_input(
            "API Key",
            type="password",
            value=st.session_state.get("api_key", ""),
            placeholder="sk-...",
            help="Your API key",
        )
        base_url = st.text_input(
            "Base URL",
            value=st.session_state.get(
                "base_url", "https://api.siliconflow.cn/v1"
            ),
            placeholder="https://api.siliconflow.cn/v1",
            help="API endpoint (OpenAI, DeepSeek, local deployment, etc.)",
        )
        model = st.text_input(
            "Model Name",
            value=st.session_state.get("model", "deepseek-ai/DeepSeek-V4-Pro"),
            placeholder="deepseek-ai/DeepSeek-V4-Pro",
            help="e.g. deepseek-ai/DeepSeek-V4-Pro",
        )

        st.session_state["api_key"] = api_key
        st.session_state["base_url"] = base_url
        st.session_state["model"] = model

    st.markdown("---")
    st.caption("ACTA v1.0  ·  Research Assistant")


# ==================== Page 5: Journal Recommender ====================
render_web_pet()


if page == "\U0001f4cb Journal Recommender":
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown(
        '<div class="page-hero">\U0001f4cb Journal Recommender</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="page-tagline">Browse and filter journals by discipline, impact factor, and zone.</div>',
        unsafe_allow_html=True,
    )

    abstract = st.text_area(
        "Paper Abstract / Topic Description",
        placeholder="Paste your paper abstract or describe your research topic. ACTA will analyze it and suggest relevant fields for journal selection.",
        height=150,
        key="journal_abstract_input",
    )

    col_left, col_right = st.columns([1, 2])
    with col_left:
        analyze_btn = st.button("Analyze", type="primary", use_container_width=True, key="journal_analyze_btn")

    if analyze_btn:
        if not abstract.strip():
            st.warning("Please enter a paper abstract.")
        else:
            try:
                with st.spinner("ACTA is analyzing your research ..."):
                    journals_all = st.session_state.get("cached_journals", [])
                    if not journals_all:
                        journals_all = load_journal_data()
                        st.session_state["cached_journals"] = journals_all
                    analysis = analyze_abstract_for_journals(abstract.strip())
                    recommendations = recommend_journals_from_abstract(
                        journals_all,
                        analysis,
                        abstract.strip(),
                        per_zone=10,
                    )
                    st.session_state["page5_analysis"] = analysis
                    st.session_state["page5_recommendations"] = recommendations
                    st.session_state["page5_abstract"] = abstract.strip()
            except ValueError as e:
                st.error("Configuration Error: " + str(e))
            except Exception as e:
                st.error("Error: " + str(e))

    if "page5_analysis" in st.session_state:
        analysis = st.session_state["page5_analysis"]
        st.markdown("---")
        st.markdown("### Field Analysis")
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown("**Broad research field:** " + ", ".join(analysis.get("broad_fields", [])))
        with col_b:
            st.markdown("**Research directions:** " + ", ".join(analysis.get("major_directions", [])[:8]))
        if analysis.get("directions"):
            st.markdown("**Semantic directions:** " + ", ".join(analysis.get("directions", [])[:12]))
        if analysis.get("reasoning"):
            st.info(analysis.get("reasoning"))

    if "page5_recommendations" in st.session_state:
        st.markdown("---")
        st.markdown("### Recommended Journals")
        st.caption("The recommender first keeps New Sharp Zones 1/2/3, then matches field, major direction, journal_intro, and expanded semantic terms.")
        render_page5_recommendations(st.session_state["page5_recommendations"])

    st.markdown("---")

    render_journal_results()

# ==================== Page 1: Scientific Question Synthesis ====================
if page == "🔬 Scientific Question Synthesis":
    # Prevent the Streamlit toolbar from overlapping the hero title
    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        '<div class="page-hero">🔬 Scientific Question Synthesis</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="page-tagline">'
        "Enter research keywords (separated by commas or semicolons) and let AI generate forward-looking "
        "scientific questions with associated and extended keyword sets."
        "</div>",
        unsafe_allow_html=True,
    )

    # Input area
    keywords = st.text_input(
        "Research Keywords",
        placeholder="e.g. quantum computing, error correction",
        help="Enter 1–5 keywords, separated by commas",
        key="keywords_input",
    )

    col_left, col_right = st.columns([1, 2])
    with col_left:
        generate_btn = st.button(
            "Generate",
            type="primary",
            use_container_width=True,
        )

    st.markdown("<hr>", unsafe_allow_html=True)

    if generate_btn:
        if not keywords.strip():
            st.warning("⚠️ Please enter at least one research keyword.")
        else:
            try:
                with st.spinner("ACTA is thinking ..."):
                    prompt_template = load_prompt_from_docx(
                        PROMPT_FILE_PAGE1, DEFAULT_PROMPT_PAGE1
                    )
                    final_prompt = prompt_template.replace(
                        "{keywords}", keywords.strip()
                    )

                    raw_result = call_llm(final_prompt)
                    questions = parse_questions(raw_result)
                    st.session_state["page1_questions"] = questions
                    st.session_state["page1_keywords"] = keywords.strip()

                if not questions:
                    st.warning(
                        "⚠️ The model did not return questions in the "
                        "expected format. Raw response shown below."
                    )
                    st.text(raw_result)
                else:
                    st.markdown(
                        "### 📝 Synthesized Questions "
                        f"({len(questions)} results)"
                    )

                    for i, q in enumerate(questions, 1):
                        ex_tags = "".join(
                            f'<span class="kw-tag extended">{k}</span>'
                            for k in q["extended"]
                        )

                        st.markdown(
                            f'<div class="question-block">'
                            f'<div class="q-title">'
                            f"<strong>Q{i}:</strong> {q['question']}"
                            f"</div>"
                            f'<div class="q-keywords">'
                            f"<strong>Extended Keywords:</strong> {ex_tags}"
                            f"</div>"
                            f"</div>",
                            unsafe_allow_html=True,
                        )

                    st.success(
                        f"✅ Successfully synthesized "
                        f"{len(questions)} research questions!"
                    )

            except ValueError as e:
                st.error(f"❌ Configuration Error: {e}")
            except Exception as e:
                st.error(
                    f"❌ API Call Failed: {e}\n\n"
                    "Troubleshooting:\n"
                    "- Verify your API Key is correct\n"
                    "- Check that the Base URL is reachable\n"
                    "- Confirm the model name exists\n"
                    "- Ensure network connectivity is stable"
                )

    # Show cached results when returning to this page
    if "page1_questions" in st.session_state and not generate_btn:
        questions = st.session_state["page1_questions"]
        if questions:
            st.markdown("### Synthesized Questions (cached) " + str(len(questions)) + " results")
            for i, q in enumerate(questions, 1):
                ex_tags = "".join("<span class=kw-tag extended>" + k + "</span>" for k in q["extended"])
                st.markdown(
                    "<div class=question-block>" +
                    "<div class=q-title><strong>Q" + str(i) + ":</strong> " + q["question"] + "</div>" +
                    "<div class=q-keywords><strong>Extended Keywords:</strong> " + ex_tags + "</div>" +
                    "</div>",
                    unsafe_allow_html=True,
                )



# ==================== Page 3: Introduction Generator ====================
elif page == "📝 Introduction Generator":
    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        '<div class="page-hero">📝 Introduction Generator</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="page-tagline">'
        "Describe your research topic and AI will generate a well-structured "
        "academic Introduction section with background, literature review, "
        "limitations, and contributions."
        "</div>",
        unsafe_allow_html=True,
    )

    # Input area
    topic_description = st.text_area(
        "Research Topic Description",
        placeholder=(
            "Describe your paper’s research topic, key ideas, methodology, "
            "and contributions in detail. The more context you provide, "
            "the better the generated Introduction will be.\n\n"
            "e.g. This paper proposes a novel deep learning framework for "
            "real-time traffic flow prediction using graph neural networks "
            "and attention mechanisms..."
        ),
        help="Provide a comprehensive description of your research topic (max ~300 words)",
        max_chars=4000,
        height=200,
        key="topic_description_input",
    )

    col_left, col_right = st.columns([1, 2])
    with col_left:
        generate_btn = st.button(
            "Generate",
            type="primary",
            use_container_width=True,
            key="generate_intro_btn",
        )

    st.markdown("<hr>", unsafe_allow_html=True)

    if generate_btn:
        if not topic_description.strip():
            st.warning("⚠️ Please enter a research topic description.")
        else:
            try:
                with st.spinner("ACTA is writing the Introduction ..."):
                    prompt_template = load_prompt_from_docx(
                        PROMPT_FILE_PAGE3, DEFAULT_PROMPT_PAGE3
                    )
                    final_prompt = prompt_template.replace(
                        "{topic}", topic_description.strip()
                    )

                    raw_result = call_llm(final_prompt)
                    st.session_state["page3_intro"] = raw_result


                    st.markdown("### 📄 Generated Introduction")
                    # Render the introduction as Markdown
                    st.markdown(raw_result)
                    st.success("✅ Introduction generated successfully!")

            except ValueError as e:
                st.error(f"❌ Configuration Error: {e}")
            except Exception as e:
                st.error(
                    f"❌ API Call Failed: {e}\n\n"
                    "Troubleshooting:\n"
                    "- Verify your API Key is correct\n"
                    "- Check that the Base URL is reachable\n"
                    "- Confirm the model name exists\n"
                    "- Ensure network connectivity is stable"
                )

    # Show cached results when returning to this page
    if "page3_intro" in st.session_state and not generate_btn:
        st.markdown("### Generated Introduction (cached)")
        st.markdown(st.session_state["page3_intro"])




# ==================== Page 4: About ACTA ====================
elif page == "ℹ️ About ACTA":
    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    # Load content from the about Word document or Markdown file
    about_content = load_prompt_from_docx(
        PROMPT_FILE_PAGE4, DEFAULT_PROMPT_PAGE4
    )

    # Render as Markdown for proper formatting
    st.markdown(about_content)

    st.markdown("<hr>", unsafe_allow_html=True)

    # Display QR code image if present in project directory
    qr_paths = []
    for fname in ["QRcode.png", "QRcode.jpg"]:
        candidate = os.path.join(APP_DIR, fname)
        if os.path.exists(candidate):
            qr_paths.append(candidate)

    if qr_paths:
        st.markdown(
            '<p style="text-align:center; color:#64748b; '
            'font-size:13px; margin-top:32px;">— — —</p>',
            unsafe_allow_html=True,
        )
        for qr_path in qr_paths:
            st.image(qr_path, width=180)
    else:
        st.markdown(
            '<p style="text-align:center; color:#94a3b8; '
            'font-size:12px; margin-top:24px;">'
            '(Place QRcode.png or QRcode.jpg in the project folder to display here)'
            '</p>',
            unsafe_allow_html=True,
        )



# ==================== Page 2: Framework Architect ====================
elif page == "📐 Framework Architect":
    # Prevent the Streamlit toolbar from overlapping the hero title
    st.markdown("<br>", unsafe_allow_html=True)

    st.markdown(
        '<div class="page-hero">'
        "📐 Framework Architect</div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="page-tagline">'
        "Provide your paper details and AI will generate a rigorous, "
        "hierarchically structured academic outline."
        "</div>",
        unsafe_allow_html=True,
    )

    # Input area
    paper_title = st.text_input(
        "Paper Title",
        placeholder="e.g. Multimodal Sentiment Analysis with Deep Learning",
        help="Enter the full title of your paper",
        key="paper_title_input",
    )

    paper_description = st.text_area(
        "Brief Description",
        placeholder=(
            "Provide up to 150 words describing your research scope, "
            "methodology, key contributions, or specific directions "
            "you want the outline to emphasize…"
        ),
        help="Optional context to guide the outline generation (max ~150 words)",
        max_chars=1200,
        height=100,
        key="paper_description_input",
    )

    paper_type = st.selectbox(
        "Paper Type",
        [
            "Modeling",
            "Experimental",
            "Survey",
        ],
        help="Select the research methodology category",
        key="paper_type_select",
    )

    col_left, col_right = st.columns([1, 2])
    with col_left:
        generate_btn = st.button(
            "Generate",
            type="primary",
            use_container_width=True,
        )

    st.markdown("<hr>", unsafe_allow_html=True)

    if generate_btn:
        if not paper_title.strip():
            st.warning("⚠️ Please enter a paper title.")
        else:
            try:
                with st.spinner("ACTA is designing the framework ..."):
                    prompt_template = load_prompt_from_docx(
                        PROMPT_FILE_PAGE2, DEFAULT_PROMPT_PAGE2
                    )
                    desc = paper_description.strip() or (
                        "No additional description provided."
                    )
                    final_prompt = (
                        prompt_template
                        .replace("{title}", paper_title.strip())
                        .replace("{paper_type}", paper_type)
                        .replace("{description}", desc)
                    )

                    raw_result = call_llm(final_prompt)
                    rendered_html = render_framework_markdown(raw_result)
                    st.session_state["page2_framework"] = rendered_html


                    st.markdown("### 📋 Paper Outline")
                    st.markdown(rendered_html, unsafe_allow_html=True)
                    st.success("✅ Framework generated successfully!")

            except ValueError as e:
                st.error(f"❌ Configuration Error: {e}")
            except Exception as e:
                st.error(
                    f"❌ API Call Failed: {e}\n\n"
                    "Troubleshooting:\n"
                    "- Verify your API Key is correct\n"
                    "- Check that the Base URL is reachable\n"
                    "- Confirm the model name exists\n"
                    "- Ensure network connectivity is stable"
                )

    # Show cached results when returning to this page
    if "page2_framework" in st.session_state and not generate_btn:
        st.markdown("### Paper Outline (cached)")
        st.markdown(st.session_state["page2_framework"], unsafe_allow_html=True)
